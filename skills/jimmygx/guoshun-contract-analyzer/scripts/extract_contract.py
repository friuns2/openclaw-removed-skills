#!/usr/bin/env python3
"""
合同文本提取脚本 - 国顺合同审核技能
支持 PDF 和 DOCX 格式
"""

import sys
import os

def extract_pdf(pdf_path):
    """从PDF提取文本"""
    try:
        import pdfplumber
    except ImportError:
        print("[ERROR] pdfplumber not installed. Install with: pip install pdfplumber", file=sys.stderr)
        return None
    
    text = ""
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
            
            # 提取表格
            tables = page.extract_tables()
            for table in tables:
                for row in table:
                    row_text = " | ".join([str(cell) if cell else "" for cell in row])
                    text += row_text + "\n"
                text += "\n"
    
    return text


def extract_docx(docx_path):
    """从DOCX提取文本"""
    try:
        from docx import Document
    except ImportError:
        print("[ERROR] python-docx not installed. Install with: pip install python-docx", file=sys.stderr)
        return None
    
    doc = Document(docx_path)
    text = ""
    
    for para in doc.paragraphs:
        text += para.text + "\n"
    
    # 提取表格
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join([cell.text for cell in row.cells])
            text += row_text + "\n"
        text += "\n"
    
    return text


def main():
    if len(sys.argv) < 2:
        print("Usage: python extract_contract.py <file.pdf|file.docx>")
        sys.exit(1)
    
    file_path = sys.argv[1]
    
    if not os.path.exists(file_path):
        print(f"[ERROR] File not found: {file_path}", file=sys.stderr)
        sys.exit(1)
    
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == ".pdf":
        text = extract_pdf(file_path)
    elif ext in [".docx", ".doc"]:
        text = extract_docx(file_path)
    else:
        print(f"[ERROR] Unsupported file type: {ext}", file=sys.stderr)
        print("Supported types: .pdf, .docx, .doc")
        sys.exit(1)
    
    if text:
        print(text)
    else:
        print("[ERROR] Failed to extract text", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
