#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文档格式验证器
验证文档是否符合GB/T 9704-2012标准和其他公文格式规范
"""

import argparse
import os
import sys
from pathlib import Path
from typing import Dict, List, Any, Optional

# 添加技能目录到路径
skill_dir = Path(__file__).parent.parent
sys.path.insert(0, str(skill_dir))

try:
    from docx import Document
    from docx.shared import Cm, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("错误：需要安装 python-docx 库")
    print("请运行：pip install python-docx")
    sys.exit(1)


class DocumentFormatValidator:
    """文档格式验证器"""
    
    def __init__(self):
        self.document = None
        self.standards = {
            'gb_t_9704_2012': self._load_gb_t_9704_2012_standard(),
            'party_government': self._load_party_government_standard(),
            'enterprise': self._load_enterprise_standard()
        }
        self.validation_results = []
    
    def validate_document(self, file_path: str, standard: str = 'gb_t_9704_2012') -> Dict[str, Any]:
        """验证文档格式
        
        Args:
            file_path: 文档文件路径
            standard: 标准类型 (gb_t_9704_2012, party_government, enterprise)
            
        Returns:
            Dict: 验证结果
        """
        if not os.path.exists(file_path):
            return {
                'valid': False,
                'errors': [f"文件不存在：{file_path}"],
                'warnings': [],
                'details': {}
            }
        
        try:
            self.document = Document(file_path)
        except Exception as e:
            return {
                'valid': False,
                'errors': [f"无法读取文档：{e}"],
                'warnings': [],
                'details': {}
            }
        
        # 选择验证标准
        if standard not in self.standards:
            return {
                'valid': False,
                'errors': [f"不支持的标准：{standard}"],
                'warnings': [],
                'details': {}
            }
        
        standard_config = self.standards[standard]
        self.validation_results = []
        
        # 执行验证
        self._validate_page_format(standard_config)
        self._validate_font_format(standard_config)
        self._validate_paragraph_format(standard_config)
        self._validate_structure(standard_config)
        
        # 统计结果
        errors = [r for r in self.validation_results if r['level'] == 'error']
        warnings = [r for r in self.validation_results if r['level'] == 'warning']
        
        return {
            'valid': len(errors) == 0,
            'errors': [e['message'] for e in errors],
            'warnings': [w['message'] for w in warnings],
            'details': {
                'total_checks': len(self.validation_results),
                'passed_checks': len(self.validation_results) - len(errors) - len(warnings),
                'failed_checks': len(errors),
                'warning_checks': len(warnings)
            }
        }
    
    def _load_gb_t_9704_2012_standard(self) -> Dict[str, Any]:
        """加载GB/T 9704-2012标准配置"""
        return {
            'name': 'GB/T 9704-2012 党政机关公文格式',
            'page_margins': {
                'top': Cm(3.7),
                'bottom': Cm(3.5),
                'left': Cm(2.8),
                'right': Cm(2.6)
            },
            'font_sizes': {
                'title': Pt(22),  # 二号
                'heading': Pt(16),  # 三号
                'body': Pt(16)  # 三号
            },
            'line_spacing': Pt(28),  # 28磅
            'first_line_indent': Cm(0.85),  # 2字符缩进
            'required_elements': [
                'title',
                'body'
            ]
        }
    
    def _load_party_government_standard(self) -> Dict[str, Any]:
        """加载党政机关公文标准配置"""
        return {
            'name': '党政机关公文格式',
            'page_margins': {
                'top': Cm(3.7),
                'bottom': Cm(3.5),
                'left': Cm(2.8),
                'right': Cm(2.6)
            },
            'font_sizes': {
                'title': Pt(22),
                'heading': Pt(16),
                'body': Pt(16)
            },
            'line_spacing': Pt(28),
            'first_line_indent': Cm(0.85),
            'required_elements': [
                'title',
                'issuing_authority',
                'document_number',
                'body'
            ]
        }
    
    def _load_enterprise_standard(self) -> Dict[str, Any]:
        """加载企业公文标准配置"""
        return {
            'name': '企业公文规范',
            'page_margins': {
                'top': Cm(2.5),
                'bottom': Cm(2.5),
                'left': Cm(2.5),
                'right': Cm(2.5)
            },
            'font_sizes': {
                'title': Pt(18),  # 稍小
                'heading': Pt(14),
                'body': Pt(12)
            },
            'line_spacing': Pt(24),
            'first_line_indent': Cm(0.85),
            'required_elements': [
                'company_logo',
                'title',
                'body',
                'signature'
            ]
        }
    
    def _validate_page_format(self, standard: Dict[str, Any]):
        """验证页面格式"""
        if not self.document or not self.document.sections:
            self._add_result('error', '文档没有页面设置')
            return
        
        section = self.document.sections[0]
        
        # 检查页边距
        margins = standard.get('page_margins', {})
        if margins:
            tolerance = Cm(0.1)  # 允许的误差范围
            
            if 'top' in margins:
                if abs(section.top_margin - margins['top']) > tolerance:
                    self._add_result('error', 
                        f"上边距不符合标准：{section.top_margin.cm:.1f}cm (应为{margins['top'].cm:.1f}cm)")
            
            if 'bottom' in margins:
                if abs(section.bottom_margin - margins['bottom']) > tolerance:
                    self._add_result('error',
                        f"下边距不符合标准：{section.bottom_margin.cm:.1f}cm (应为{margins['bottom'].cm:.1f}cm)")
            
            if 'left' in margins:
                if abs(section.left_margin - margins['left']) > tolerance:
                    self._add_result('error',
                        f"左边距不符合标准：{section.left_margin.cm:.1f}cm (应为{margins['left'].cm:.1f}cm)")
            
            if 'right' in margins:
                if abs(section.right_margin - margins['right']) > tolerance:
                    self._add_result('error',
                        f"右边距不符合标准：{section.right_margin.cm:.1f}cm (应为{margins['right'].cm:.1f}cm)")
        
        # 检查纸张大小
        expected_width = Cm(21.0)  # A4宽度
        expected_height = Cm(29.7)  # A4高度
        tolerance = Cm(0.1)
        
        if abs(section.page_width - expected_width) > tolerance:
            self._add_result('warning',
                f"纸张宽度不符合A4标准：{section.page_width.cm:.1f}cm (应为21.0cm)")
        
        if abs(section.page_height - expected_height) > tolerance:
            self._add_result('warning',
                f"纸张高度不符合A4标准：{section.page_height.cm:.1f}cm (应为29.7cm)")
    
    def _validate_font_format(self, standard: Dict[str, Any]):
        """验证字体格式"""
        if not self.document or not self.document.paragraphs:
            self._add_result('error', '文档没有内容')
            return
        
        font_sizes = standard.get('font_sizes', {})
        
        # 检查标题字体
        if 'title' in font_sizes:
            title_paragraphs = [p for p in self.document.paragraphs if p.text.strip()]
            if title_paragraphs:
                first_para = title_paragraphs[0]
                if first_para.runs:
                    first_run = first_para.runs[0]
                    if first_run.font.size:
                        if abs(first_run.font.size.pt - font_sizes['title'].pt) > 1:
                            self._add_result('warning',
                                f"标题字体大小不符合标准：{first_run.font.size.pt:.1f}pt (应为{font_sizes['title'].pt:.1f}pt)")
        
        # 检查正文字体
        body_font_size = font_sizes.get('body', Pt(16))
        for i, paragraph in enumerate(self.document.paragraphs[1:5]):  # 检查前几个正文段落
            if paragraph.runs:
                for run in paragraph.runs:
                    if run.font.size:
                        if abs(run.font.size.pt - body_font_size.pt) > 1:
                            self._add_result('warning',
                                f"第{i+2}段正文字体大小不符合标准：{run.font.size.pt:.1f}pt (应为{body_font_size.pt:.1f}pt)")
    
    def _validate_paragraph_format(self, standard: Dict[str, Any]):
        """验证段落格式"""
        if not self.document or not self.document.paragraphs:
            return
        
        # 检查行距
        expected_line_spacing = standard.get('line_spacing', Pt(28))
        tolerance = Pt(2)
        
        for i, paragraph in enumerate(self.document.paragraphs[:10]):  # 检查前10个段落
            if paragraph.paragraph_format.line_spacing:
                line_spacing_pt = paragraph.paragraph_format.line_spacing.pt if hasattr(paragraph.paragraph_format.line_spacing, 'pt') else 0
                if abs(line_spacing_pt - expected_line_spacing.pt) > tolerance.pt:
                    self._add_result('warning',
                        f"第{i+1}段行距不符合标准：{line_spacing_pt:.1f}pt (应为{expected_line_spacing.pt:.1f}pt)")
        
        # 检查首行缩进
        expected_indent = standard.get('first_line_indent', Cm(0.85))
        tolerance = Cm(0.1)
        
        for i, paragraph in enumerate(self.document.paragraphs[1:6]):  # 检查第2-6个段落（正文段落）
            if paragraph.paragraph_format.first_line_indent:
                indent = paragraph.paragraph_format.first_line_indent
                indent_cm = indent.cm if hasattr(indent, 'cm') else 0
                if abs(indent_cm - expected_indent.cm) > tolerance.cm:
                    self._add_result('warning',
                        f"第{i+2}段首行缩进不符合标准：{indent_cm:.1f}cm (应为{expected_indent.cm:.1f}cm)")
    
    def _validate_structure(self, standard: Dict[str, Any]):
        """验证文档结构"""
        if not self.document or not self.document.paragraphs:
            self._add_result('error', '文档没有内容')
            return
        
        # 检查必需元素
        required_elements = standard.get('required_elements', [])
        
        paragraphs_text = [p.text.strip() for p in self.document.paragraphs if p.text.strip()]
        
        if 'title' in required_elements:
            if not paragraphs_text:
                self._add_result('error', '文档缺少标题')
            elif len(paragraphs_text[0]) < 2:  # 标题太短
                self._add_result('warning', '文档标题可能不完整')
        
        if 'body' in required_elements:
            if len(paragraphs_text) < 3:  # 至少标题+2段正文
                self._add_result('warning', '文档正文内容可能不足')
        
        # 检查标题对齐
        if paragraphs_text:
            first_para = self.document.paragraphs[0]
            if first_para.paragraph_format.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                self._add_result('warning', '文档标题未居中')
    
    def _add_result(self, level: str, message: str):
        """添加验证结果"""
        self.validation_results.append({
            'level': level,
            'message': message
        })


def print_validation_results(results: Dict[str, Any], detailed: bool = False):
    """打印验证结果"""
    print("\n" + "="*60)
    print(f"文档格式验证报告")
    print("="*60)
    
    if results['valid']:
        print("✅ 文档格式验证通过！")
    else:
        print("❌ 文档格式验证未通过！")
    
    print(f"\n📊 验证统计：")
    details = results['details']
    print(f"   总检查项：{details.get('total_checks', 0)}")
    print(f"   通过项：{details.get('passed_checks', 0)}")
    print(f"   失败项：{details.get('failed_checks', 0)}")
    print(f"   警告项：{details.get('warning_checks', 0)}")
    
    if results['errors']:
        print(f"\n❌ 错误（{len(results['errors'])}个）：")
        for i, error in enumerate(results['errors'], 1):
            print(f"   {i}. {error}")
    
    if results['warnings']:
        print(f"\n⚠️ 警告（{len(results['warnings'])}个）：")
        for i, warning in enumerate(results['warnings'], 1):
            print(f"   {i}. {warning}")
    
    if detailed:
        print(f"\n📋 详细建议：")
        if results['errors']:
            print("   1. 请先解决所有错误项")
        if results['warnings']:
            print("   2. 建议处理所有警告项以提升文档质量")
        
        print("   3. 确保文档符合相关标准格式要求")
        print("   4. 使用文档模板可以确保格式一致性")
    
    print("\n" + "="*60)


def main():
    parser = argparse.ArgumentParser(description='文档格式验证器')
    parser.add_argument('--file', required=True, help='要验证的文档文件路径')
    parser.add_argument('--standard', default='gb_t_9704_2012',
                       choices=['gb_t_9704_2012', 'party_government', 'enterprise'],
                       help='验证标准')
    parser.add_argument('--detailed', action='store_true', help='显示详细报告')
    parser.add_argument('--output', help='输出报告文件路径')
    
    args = parser.parse_args()
    
    # 验证文档
    validator = DocumentFormatValidator()
    results = validator.validate_document(args.file, args.standard)
    
    # 打印结果
    print_validation_results(results, args.detailed)
    
    # 保存报告
    if args.output:
        try:
            import json
            with open(args.output, 'w', encoding='utf-8') as f:
                json.dump(results, f, ensure_ascii=False, indent=2)
            print(f"\n报告已保存到：{args.output}")
        except Exception as e:
            print(f"保存报告失败：{e}")
    
    # 返回退出码
    if not results['valid']:
        sys.exit(1)
    else:
        sys.exit(0)


if __name__ == "__main__":
    main()