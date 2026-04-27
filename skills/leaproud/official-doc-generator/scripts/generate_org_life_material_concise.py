#!/usr/bin/env python3
# -*- coding: utf-8 -*-
'''
组织生活会个人对照检查材料生成器 - 精简版
目标字数 800-1000 字
'''

import os
import sys
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print('错误：需要安装 python-docx 库')
    sys.exit(1)


TITLE = '组织生活会个人对照检查材料'

PARAGRAPHS = [
    ('h1', '一、对照检查剖析'),
    ('h2', '（一）筑牢政治忠诚'),
    ('body', '在政治忠诚方面，本人坚持以习近平新时代中国特色社会主义思想为指导，坚决维护党中央权威和集中统一领导。但对照更高标准，理论学习的系统性还不够，有时存在\u201c泛泛而学\u201d现象；理论联系实际不够紧密，学用脱节的问题依然存在；主动发声亮剑的意识还需增强。'),
    ('h2', '（二）加强党性锤炼'),
    ('body', '在党性锤炼方面，本人自觉参加组织生活，落实\u201c三会一课\u201d制度。但深刻反思，党性锻炼的主动性有所欠缺，有时存在\u201c过得去\u201d的思想；宗旨意识树得不够牢固，面对繁杂事务时容易产生急躁情绪；艰苦奋斗精神有所弱化，攻坚克难的劲头还不够足。'),
    ('h2', '（三）严守纪律规矩'),
    ('body', '在纪律规矩方面，本人严格遵守各项纪律规定，自觉执行中央八项规定精神。但对照检查，纪律意识的自觉性还不够高，对\u201c小节\u201d问题不够重视；执行规章制度有时存在打折扣现象；请示报告制度落实不够到位，个别事项存在事后报告的情况。'),
    ('h2', '（四）改作风树新风'),
    ('body', '在作风建设方面，本人注重勤勉务实、廉洁自律。但对照新时代要求，工作作风还不够扎实，存在形式主义倾向，个别工作满足于\u201c做了\u201d而非\u201c做好了\u201d；调查研究不够深入，掌握第一手资料不够全面；担当精神还不够强，面对矛盾有时存在畏难情绪。'),
    ('h2', '（五）后勤领域专项整肃治理'),
    ('body', '对照后勤领域\u201c六个方面\u201d要求，一是政治忠诚方面，对后勤工作的政治属性认识不够深刻，政治把关能力有待提高；二是党性原则方面，坚持原则的定力不够，个别事项处理存在\u201c和稀泥\u201d现象；三是履职用权方面，审批流程把关不够严格，信息公开透明度有待提高；四是遵规守纪方面，廉政风险排查不够全面，物资管理精细化水平需提升；五是行业操守方面，与供应商交往的边界意识需强化，服务品质化水平有待提升；六是赓续传统方面，勤俭节约意识有所淡化，过紧日子的思想树立不够牢固。'),
    ('h1', '二、下一步整改措施'),
    ('body', '1.强化理论武装。制定个人学习计划，系统研读原著原文，注重理论联系实际，每季度撰写学习心得，切实将学习成果转化为工作能力。'),
    ('body', '2.锤炼党性修养。自觉参加党内政治生活，认真开展批评与自我批评，牢固树立服务理念，深入基层了解需求，每季度至少开展一次专项调研，解决服务对象反映的突出问题。'),
    ('body', '3.严守纪律底线。持续加强党规党纪学习，严格执行请示报告制度，杜绝漏报迟报现象；完善后勤廉政风险防控机制，全面排查风险点，建立台账逐项整改，保持清正廉洁的政治本色。'),
    ('body', '4.改进工作作风。坚决整治形式主义，大兴调查研究之风，每年不少于两次专项调研，确保决策有依据、工作有实效；强化担当意识，面对矛盾敢于碰硬，做到守土有责、守土尽责。'),
    ('body', '5.深化后勤整肃。对照\u201c六个方面\u201d逐项整改销号，完善后勤管理制度体系，规范关键环节操作流程，推进后勤信息公开，坚持勤俭办一切事业，杜绝铺张浪费。'),
]


class OrgLifeDocGenerator:

    def __init__(self):
        self.doc = Document()
        self._setup_page()
        self._setup_styles()

    def _setup_page(self):
        sec = self.doc.sections[0]
        sec.page_width = Cm(21.0)
        sec.page_height = Cm(29.7)
        sec.top_margin = Cm(3.7)
        sec.bottom_margin = Cm(3.5)
        sec.left_margin = Cm(2.8)
        sec.right_margin = Cm(2.6)
        sec.header_distance = Cm(1.5)
        sec.footer_distance = Cm(1.75)

    def _setup_styles(self):
        sn = self.doc.styles['Normal']
        sn.font.name = '\u4eff\u5b8b'
        sn.font.size = Pt(16)
        sn.paragraph_format.line_spacing = Pt(28)
        sn.paragraph_format.first_line_indent = Cm(0.85)
        rpr = sn._element.get_or_add_rPr()
        rpr.rFonts.set(qn('w:eastAsia'), '\u4eff\u5b8b')
        rpr.rFonts.set(qn('w:ascii'), 'Times New Roman')
        rpr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')

    def _set_run_font(self, run, east_asia, size_pt):
        run.font.name = east_asia
        run.font.size = Pt(size_pt)
        rpr = run._r.get_or_add_rPr()
        rpr.rFonts.set(qn('w:eastAsia'), east_asia)
        rpr.rFonts.set(qn('w:ascii'), 'Times New Roman')
        rpr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')

    def _add_title(self, text):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(28)
        p.paragraph_format.line_spacing = Pt(28)
        run = p.add_run(text)
        self._set_run_font(run, '\u65b9\u6b63\u5c0f\u6807\u5b8b\u7b80\u4f53', 22)
        return p

    def _add_h1(self, text):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = Pt(28)
        p.paragraph_format.first_line_indent = Cm(0)
        run = p.add_run(text)
        self._set_run_font(run, '\u9ed1\u4f53', 16)
        return p

    def _add_h2(self, text):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = Pt(28)
        p.paragraph_format.first_line_indent = Cm(0.85)
        run = p.add_run(text)
        self._set_run_font(run, '\u6977\u4f53', 16)
        return p

    def _add_body(self, text):
        p = self.doc.add_paragraph()
        p.paragraph_format.line_spacing = Pt(28)
        p.paragraph_format.first_line_indent = Cm(0.85)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text)
        self._set_run_font(run, '\u4eff\u5b8b', 16)
        return p

    def build(self):
        self._add_title(TITLE)
        handlers = {
            'h1': self._add_h1,
            'h2': self._add_h2,
            'body': self._add_body,
        }
        for ptype, text in PARAGRAPHS:
            handlers[ptype](text)

    def save(self, output_path):
        self.doc.save(output_path)
        print(f'文档已保存: {output_path}')
        total = sum(len(p.text) for p in self.doc.paragraphs)
        print(f'总字数: 约 {total} 字')


def main():
    import argparse
    parser = argparse.ArgumentParser(description='组织生活会个人对照检查材料生成器（精简版）')
    parser.add_argument('--output-dir', required=True,
                        help='输出目录（必须明确指定，禁止默认写入桌面或任何硬编码路径）')
    parser.add_argument('--filename', default=None, help='输出文件名（不含路径）')
    args = parser.parse_args()

    output_dir = args.output_dir
    os.makedirs(output_dir, exist_ok=True)
    date_str = datetime.now().strftime('%Y%m%d')
    if args.filename:
        output_path = os.path.join(output_dir, args.filename)
    else:
        output_path = os.path.join(output_dir, f'组织生活会个人对照检查材料_精简版_{date_str}.docx')

    gen = OrgLifeDocGenerator()
    gen.build()
    gen.save(output_path)
    print(f'完成！文件位置: {output_path}')


if __name__ == '__main__':
    main()
