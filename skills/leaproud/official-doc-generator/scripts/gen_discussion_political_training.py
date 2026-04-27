#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
学习讨论材料生成器（示例脚本）
格式规范（GB/T 9704-2012）：
- 标题：2号方正小标宋简体，居中，直接顶格
- 正文：3号仿宋_GB2312，两端对齐，首行缩进2字符
- 一级标题：3号黑体，首行缩进2字符，两端对齐
- 署名：3号楷体_GB2312，居中
- 英文：Times New Roman
- 页码：4号宋体，单页右/双页左，一字线，版心下7mm

用法：
  python gen_discussion_political_training.py --title "XXX" --author "单位 姓名" --output-dir "D:\\输出目录"
"""

import argparse
import os
import sys
from pathlib import Path

skill_dir = Path(__file__).parent.parent
sys.path.insert(0, str(skill_dir))

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ========== 默认配置（运行时由命令行参数覆盖）==========
DEFAULT_TITLE = "学习重要讲话精神的讨论"
DEFAULT_AUTHOR = ""   # 默认空署名，由用户通过 --author 指定
DEFAULT_OUTPUT_DIR = ""  # 必须由用户通过 --output-dir 指定

ENGLISH_FONT = "Times New Roman"
CHAR_INDENT = Pt(32)  # 2字符缩进（16pt × 2）


def create_document(title=None, author=None, output_dir=None):
    """
    生成文档。
    Args:
        title: 文档标题（支持 \\n 换行），默认使用 DEFAULT_TITLE
        author: 署名，格式如"单位 姓名"，为空则不添加署名行
        output_dir: 输出目录（必须明确指定）
    """
    if not output_dir:
        raise ValueError("必须通过 --output-dir 指定输出目录，不允许自动写入任何硬编码路径。")

    doc = Document()

    # ===== 页面设置 =====
    section = doc.sections[0]
    section.top_margin = Cm(3.7)
    section.bottom_margin = Cm(3.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.odd_and_even_pages_header_footer = True

    # 用原始XML强制确保 evenAndOddHeaders 标记写入（WPS兼容）
    sectPr = section._sectPr
    if sectPr is not None:
        existing = sectPr.find(qn('w:evenAndOddHeaders'))
        if existing is None:
            even_odd = OxmlElement('w:evenAndOddHeaders')
            sectPr.insert(0, even_odd)

    # ===== 页码（奇数页居右，偶数页居左，一字线） =====
    add_page_numbers(section)

    # 在 settings.xml 中强制添加 evenAndOddHeaders（WPS兼容，否则WPS不区分奇偶页脚）
    settings_el = doc.settings.element
    existing = settings_el.find(qn('w:evenAndOddHeaders'))
    if existing is None:
        even_odd = OxmlElement('w:evenAndOddHeaders')
        settings_el.append(even_odd)

    # ===== 标题（2号方正小标宋简体，居中，直接顶格） =====
    use_title = (title or DEFAULT_TITLE).replace('\\n', '\n')
    title_lines = use_title.split('\n')
    for i, line in enumerate(title_lines):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(28.95)
        # 直接顶格，不留白
        run = p.add_run(line)
        run.font.size = Pt(22)
        run.font.name = '方正小标宋简体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '方正小标宋简体')
        run._element.rPr.rFonts.set(qn('w:ascii'), ENGLISH_FONT)
        run._element.rPr.rFonts.set(qn('w:hAnsi'), ENGLISH_FONT)


    # ===== 署名（3号楷体，居中；仅在指定了 author 时输出）=====
    use_author = author or DEFAULT_AUTHOR
    if use_author:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(28.95)
        run = p.add_run(use_author)
        run.font.size = Pt(16)
        run.font.name = '楷体_GB2312'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体_GB2312')
        run._element.rPr.rFonts.set(qn('w:ascii'), ENGLISH_FONT)
        run._element.rPr.rFonts.set(qn('w:hAnsi'), ENGLISH_FONT)

    # ===== 正文 =====
    add_normal(doc,
        "通过认真学习讲话精神，结合本单位工作实际，现就如何贯彻落实讲话精神、推动思想建设与业务工作深度融合，谈几点认识和体会。")

    # 一、深刻认识重大意义
    add_h1(doc, "一、深刻认识学习教育的重大意义")

    add_normal(doc,
        "领导同志在讲话中强调，要切实加强思想政治建设，筑牢信仰之基、补足精神之钙、把稳思想之舵。这一论述高屋建瓴、切中要害，充分体现了党委对思想政治建设的高度重视。对于我们科室而言，深刻认识学习教育的重大意义，就是要准确把握以下几个层面。")

    add_h2(doc,
        "（一）从政治高度看，学习教育是坚定理想信念、筑牢政治根基的重要抓手。",
        "只有不断强化政治意识、大局意识、核心意识、看齐意识，才能在各项工作中始终保持正确的政治方向，确保党中央的决策部署在基层落地生根。")

    add_h2(doc,
        "（二）从使命任务看，学习教育是提升履职能力的现实需要。",
        "新时代对各级干部的能力素质提出了更高要求，通过持续深化学习，进一步强化使命担当，以更高的标准做好本职工作，是新时代赋予我们的神圣职责。")

    add_h2(doc,
        "（三）从自身建设看，学习教育是科室发展的根本保证。",
        "近年来，科室在业务建设、人才培养、创新发展等方面取得了一定成绩，但对照党委要求和高标准还有差距。只有把学习教育作为推动科室建设的重要契机，才能实现业务能力与综合素质的同步提升。")

    # 二、准确把握目标要求
    add_h1(doc, "二、准确把握学习的目标要求")

    add_normal(doc,
        "领导同志在讲话中明确指出，此次学习要聚焦问题导向，坚持刀刃向内，真正把自己摆进去、把职责摆进去、把工作摆进去。对照这一要求，结合科室实际，我们进行了认真的自查自纠，主要有以下几个方面需要重点改进。")

    add_h2(doc,
        "（一）理论学习还不够深入。",
        "部分同志在学习党的创新理论方面存在浅尝辄止的现象，满足于完成规定动作，缺乏深钻细研的劲头，理论联系实际不够紧密，运用科学理论指导实际工作的能力有待进一步提升。")

    add_h2(doc,
        "（二）政治敏锐性还不够强。",
        "在日常工作生活中，个别同志对社会上的错误思潮和负面言论缺乏足够的警惕性，政治鉴别力和政治免疫力还不够强，未能充分认识到新时代对党员干部的特殊政治要求。")

    add_h2(doc,
        "（三）纪律规矩意识还需加强。",
        "虽然科室整体作风良好，但在日常管理和工作中还存在一些不容忽视的薄弱环节，如请示报告制度执行不够严格、工作标准不够高等问题，需要在学习教育中认真加以解决。")

    add_h2(doc,
        "（四）服务意识还需深化。",
        "在服务群众的实践中，还存在服务方式相对单一、主动上门服务意识不够强等问题，距离群众对高质量服务的期望还有一定差距，需要以学习教育为契机加以改进提升。")

    # 三、紧密结合科室实际
    add_h1(doc, "三、结合科室实际抓落实")

    add_normal(doc,
        "领导同志强调，学习教育不能空对空，必须实打实，要与本职工作深度融合、相互促进。贯彻落实这一要求，我们将重点从以下四个方面着手。")

    add_h2(doc,
        "（一）强化理论武装，筑牢思想根基。",
        "将学习内容纳入科室常态化学习计划，坚持每周集中学习不少于一次，采取领学、研讨、交流发言相结合的方式，确保学深悟透。特别要注重将理论学习与工作实践相结合，不断提高运用科学理论解决实际问题的能力。")

    add_h2(doc,
        "（二）聚焦主责主业，提升服务质效。",
        "将学习成效直接转化为提升服务质量的实际成果。围绕核心业务任务，进一步优化工作流程，完善服务机制。加大主动服务力度，深入一线开展宣讲和指导，真正做到群众在哪里、服务就跟进到哪里。")

    add_h2(doc,
        "（三）严守纪律规矩，锤炼过硬作风。",
        "以学习教育为抓手，全面加强科室纪律作风建设。严格落实请示报告、值班执勤、保密管理等各项制度，规范工作流程，坚决杜绝任何形式的违规违纪问题。强化责任担当意识，对工作中发现的问题不回避、不遮掩，做到即知即改、立行立改。")

    add_h2(doc,
        "（四）加强队伍建设，提升整体素质。",
        "注重在学习中发现和培养骨干力量，发挥党员先锋模范作用，带动全科同志共同进步。定期开展业务培训和岗位练兵，提升专业技能水平。关心科室同志的工作生活，营造团结奋进、风清气正的良好氛围，打造一支政治过硬、业务精湛、作风优良的工作队伍。")

    # 四、立足长远谋划发展
    add_h1(doc, "四、以学习成果推动科室建设")

    add_normal(doc,
        "学习教育既是一次思想淬炼，更是一次推动工作的强大动力。我们要把学习成果转化为推动科室长远发展的实际成效，重点在以下几个方面持续用力。")

    add_h2(doc,
        "（一）对标党委要求，制定科室建设发展清单。",
        "结合学习中查摆出的问题和不足，逐条逐项制定整改措施，明确责任人和完成时限，做到有目标、有计划、有落实、有检查，确保整改到位。")

    add_h2(doc,
        "（二）着眼长远发展，深化业务研究。",
        "围绕事业发展需求，加大科研攻关力度，在关键领域和薄弱环节开展深入研究，力争产出一批有价值的研究成果，为科室可持续发展提供有力支撑。")

    add_h2(doc,
        "（三）加强协作配合，提升综合保障能力。",
        "主动加强与兄弟科室和上级业务部门的沟通协作，充分利用现有资源，形成工作合力。积极探索各项工作之间的有机融合，构建全方位、多层次的服务体系。")

    add_normal(doc,
        "总之，我们将以讲话精神为指引，以此次学习为新的起点，进一步提高政治站位，强化使命担当，以更高的标准、更严的要求、更实的举措，全力做好各项工作，为事业发展和群众利益作出新的更大贡献。")

    # ===== 保存 =====
    from datetime import datetime
    date_str = datetime.now().strftime('%Y%m%d')
    safe_title = (use_title.replace('\n', '').replace('\\n', ''))[:30]
    filename = f"{safe_title}_{date_str}.docx"
    output_file = os.path.join(output_dir, filename)
    os.makedirs(output_dir, exist_ok=True)
    doc.save(output_file)
    print(f"文档已保存到：{output_file}")
    return output_file


def add_normal(doc, text):
    """正文段落（3号仿宋，两端对齐，首行缩进2字符）"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(28.95)
    p.paragraph_format.first_line_indent = CHAR_INDENT
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    run = p.add_run(text)
    run.font.size = Pt(16)
    run.font.name = '仿宋_GB2312'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    run._element.rPr.rFonts.set(qn('w:ascii'), ENGLISH_FONT)
    run._element.rPr.rFonts.set(qn('w:hAnsi'), ENGLISH_FONT)
    return p


def add_h1(doc, text):
    """一级标题（3号黑体，首行缩进2字符，两端对齐）"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(28.95)
    p.paragraph_format.first_line_indent = CHAR_INDENT
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    run = p.add_run(text)
    run.font.size = Pt(16)
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run._element.rPr.rFonts.set(qn('w:ascii'), ENGLISH_FONT)
    run._element.rPr.rFonts.set(qn('w:hAnsi'), ENGLISH_FONT)
    return p


def add_h2(doc, heading, body=None):
    """二级标题段落（段内混排）
    GB/T 9704-2012 7.3.3: 第二层"（一）"用楷体字标注，正文续写用仿宋
    heading: 楷体标题句（如"（一）从政治高度看，学习教育是坚定理想信念的重要抓手。"）
    body: 仿宋正文部分（可选，紧接标题句后面，不换段）
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(28.95)
    p.paragraph_format.first_line_indent = CHAR_INDENT
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # 标题部分用楷体
    rh = p.add_run(heading)
    rh.font.size = Pt(16)
    rh.font.name = '楷体_GB2312'
    rh._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体_GB2312')
    rh._element.rPr.rFonts.set(qn('w:ascii'), ENGLISH_FONT)
    rh._element.rPr.rFonts.set(qn('w:hAnsi'), ENGLISH_FONT)

    # 正文部分用仿宋（如果有的话）
    if body:
        rb = p.add_run(body)
        rb.font.size = Pt(16)
        rb.font.name = '仿宋_GB2312'
        rb._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        rb._element.rPr.rFonts.set(qn('w:ascii'), ENGLISH_FONT)
        rb._element.rPr.rFonts.set(qn('w:hAnsi'), ENGLISH_FONT)

    return p


def add_page_numbers(section):
    """添加标准页码（4号宋体，单页右/双页左，一字线）"""
    try:
        def _setup_page_number(paragraph, alignment):
            """在一个页脚段落中设置一字线 + PAGE 域"""
            paragraph.text = ""
            paragraph.alignment = alignment
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            paragraph.paragraph_format.line_spacing = Pt(14)

            # 左一字线
            r1 = paragraph.add_run("— ")
            r1.font.name = '宋体'
            r1.font.size = Pt(14)
            r1._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

            # PAGE 域
            rn = paragraph.add_run()
            rn.font.name = '宋体'
            rn.font.size = Pt(14)
            rn._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            fldBegin = OxmlElement('w:fldChar')
            fldBegin.set(qn('w:fldCharType'), 'begin')
            instr = OxmlElement('w:instrText')
            instr.set(qn('xml:space'), 'preserve')
            instr.text = " PAGE "
            fldEnd = OxmlElement('w:fldChar')
            fldEnd.set(qn('w:fldCharType'), 'end')
            rn._r.append(fldBegin)
            rn._r.append(instr)
            rn._r.append(fldEnd)

            # 右一字线
            r2 = paragraph.add_run(" —")
            r2.font.name = '宋体'
            r2.font.size = Pt(14)
            r2._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # --- 奇数页页脚（居右空一字）---
        odd_footer = section.footer
        odd_footer.is_linked_to_previous = False
        _setup_page_number(odd_footer.paragraphs[0], WD_ALIGN_PARAGRAPH.RIGHT)

        # --- 偶数页页脚（居左空一字）---
        even_footer = section.even_page_footer
        even_footer.is_linked_to_previous = False
        _setup_page_number(even_footer.paragraphs[0], WD_ALIGN_PARAGRAPH.LEFT)

    except Exception as e:
        print(f"页码设置失败：{e}")


if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser(description='学习讨论材料生成器')
    parser.add_argument('--title', default=DEFAULT_TITLE, help='文档标题（可用 \\n 换行）')
    parser.add_argument('--author', default='', help='署名，格式如"单位 姓名"，留空则不署名')
    parser.add_argument('--output-dir', required=True, help='输出目录（必须指定）')
    args = parser.parse_args()
    create_document(title=args.title, author=args.author, output_dir=args.output_dir)

