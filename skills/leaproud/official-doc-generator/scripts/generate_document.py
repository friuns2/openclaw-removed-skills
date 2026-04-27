#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
官方文档生成器 - 主脚本
支持会议纪要、发言稿、讨论提纲、工作汇报等多种文档类型
遵循GB/T 9704-2012标准格式
"""

import argparse
import json
import os
import sys
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, Optional

# 添加技能目录到路径
skill_dir = Path(__file__).parent.parent
sys.path.insert(0, str(skill_dir))

try:
    from docx import Document
    from docx.shared import Pt, Cm, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    print("错误：需要安装 python-docx 库")
    print("请运行：pip install python-docx")
    sys.exit(1)


class OfficialDocumentGenerator:
    """官方文档生成器"""
    
    def __init__(self):
        self.document = None
        self.styles = {}
        
    def create_document(self, doc_type: str, config: Dict[str, Any]) -> Document:
        """创建文档
        
        Args:
            doc_type: 文档类型 (meeting_minutes, speech, discussion_outline, work_report)
            config: 配置参数
            
        Returns:
            Document: Word文档对象
        """
        # 创建新文档
        self.document = Document()
        
        # 设置页面格式
        self._setup_page_format()
        
        # 设置样式
        self._setup_styles()
        
        # 根据文档类型生成内容
        if doc_type == "meeting_minutes":
            self._generate_meeting_minutes(config)
        elif doc_type == "speech":
            self._generate_speech(config)
        elif doc_type == "discussion_outline":
            self._generate_discussion_outline(config)
        elif doc_type == "work_report":
            self._generate_work_report(config)
        else:
            raise ValueError(f"不支持的文档类型: {doc_type}")
            
        return self.document
    
    def _setup_page_format(self):
        """设置页面格式（GB/T 9704-2012标准）"""
        sections = self.document.sections
        for section in sections:
            # 设置页边距（单位：厘米）
            section.top_margin = Cm(3.7)
            section.bottom_margin = Cm(3.5)
            section.left_margin = Cm(2.8)
            section.right_margin = Cm(2.6)
            
            # 设置纸张大小
            section.page_width = Cm(21.0)  # A4宽度
            section.page_height = Cm(29.7)  # A4高度
            
    def _setup_styles(self):
        """设置文档样式"""
        # 标题样式（二号小标宋）
        try:
            title_style = self.document.styles.add_style('Title_Style', WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.name = '小标宋'
            title_style.font.size = Pt(22)  # 二号
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_after = Pt(28)
        except:
            # 如果小标宋字体不存在，使用黑体
            title_style = self.document.styles.add_style('Title_Style', WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.name = '黑体'
            title_style.font.size = Pt(22)
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_after = Pt(28)
        
        # 一级标题样式（黑体）
        h1_style = self.document.styles.add_style('Heading1_Style', WD_STYLE_TYPE.PARAGRAPH)
        h1_style.font.name = '黑体'
        h1_style.font.size = Pt(16)  # 三号
        h1_style.paragraph_format.left_indent = Cm(0)
        h1_style.paragraph_format.space_before = Pt(20)
        h1_style.paragraph_format.space_after = Pt(10)
        
        # 二级标题样式（楷体）
        h2_style = self.document.styles.add_style('Heading2_Style', WD_STYLE_TYPE.PARAGRAPH)
        h2_style.font.name = '楷体'
        h2_style.font.size = Pt(16)
        h2_style.paragraph_format.left_indent = Cm(0.85)  # 缩进0.85厘米
        h2_style.paragraph_format.space_before = Pt(10)
        h2_style.paragraph_format.space_after = Pt(5)
        
        # 正文样式（仿宋GB2312）
        try:
            normal_style = self.document.styles['Normal']
            normal_style.font.name = '仿宋 GB2312'
        except:
            normal_style = self.document.styles['Normal']
            normal_style.font.name = '仿宋'  # 备用字体
        
        normal_style.font.size = Pt(16)  # 三号
        normal_style.paragraph_format.line_spacing = Pt(28)  # 28磅行距
        normal_style.paragraph_format.first_line_indent = Cm(0.85)  # 首行缩进2字符
        
        # 保存样式
        self.styles = {
            'title': title_style,
            'h1': h1_style,
            'h2': h2_style,
            'normal': normal_style
        }
    
    def _generate_meeting_minutes(self, config: Dict[str, Any]):
        """生成会议纪要"""
        # 标题
        title = config.get('title', '会议纪要')
        self._add_paragraph(f"{title}会议纪要", style='title')
        
        # 会议基本情况
        self._add_heading("一、会议基本情况", level=1)
        
        meeting_info = config.get('meeting_info', {})
        self._add_paragraph(f"（一）时间：{meeting_info.get('time', '')}", style='normal')
        self._add_paragraph(f"（二）地点：{meeting_info.get('location', '')}", style='normal')
        self._add_paragraph(f"（三）主持人：{meeting_info.get('host', '')}", style='normal')
        self._add_paragraph(f"（四）记录人：{meeting_info.get('recorder', '')}", style='normal')
        
        # 参会人员
        attendees = meeting_info.get('attendees', [])
        if attendees:
            attendees_str = "、".join(attendees)
            self._add_paragraph(f"（五）出席人员：{attendees_str}", style='normal')
        
        absentees = meeting_info.get('absentees', [])
        if absentees:
            absentees_str = "、".join(absentees)
            self._add_paragraph(f"（六）缺席人员：{absentees_str}", style='normal')
        
        # 主要议题
        topics = config.get('topics', [])
        if topics:
            self._add_heading("二、会议议题", level=1)
            for i, topic in enumerate(topics, 1):
                self._add_paragraph(f"{i}. {topic}", style='normal')
        
        # 会议内容
        contents = config.get('contents', [])
        if contents:
            self._add_heading("三、会议主要内容", level=1)
            for i, content in enumerate(contents, 1):
                self._add_heading(f"（一）{content.get('topic', '')}", level=2)
                points = content.get('points', [])
                for j, point in enumerate(points, 1):
                    self._add_paragraph(f"{j}. {point}", style='normal')
        
        # 决议事项
        resolutions = config.get('resolutions', [])
        if resolutions:
            self._add_heading("四、会议决议事项", level=1)
            for i, resolution in enumerate(resolutions, 1):
                self._add_paragraph(f"（一）{resolution.get('content', '')}", style='normal')
                self._add_paragraph(f"    责任部门/责任人：{resolution.get('responsible', '')}", style='normal')
                self._add_paragraph(f"    完成时限：{resolution.get('deadline', '')}", style='normal')
        
        # 下一步工作安排
        next_steps = config.get('next_steps', [])
        if next_steps:
            self._add_heading("五、下一步工作安排", level=1)
            for i, step in enumerate(next_steps, 1):
                self._add_paragraph(f"{i}. {step}", style='normal')
        
        # 其他事项
        other_matters = config.get('other_matters', [])
        if other_matters:
            self._add_heading("六、其他事项", level=1)
            for i, matter in enumerate(other_matters, 1):
                self._add_paragraph(f"{i}. {matter}", style='normal')
        
        # 落款
        self._add_paragraph("", style='normal')
        self._add_paragraph("", style='normal')
        self._add_paragraph(f"主持人签字：", style='normal')
        self._add_paragraph(f"记录人签字：", style='normal')
        
        # 日期
        current_date = datetime.now().strftime("%Y年%m月%d日")
        self._add_paragraph(f"{current_date}", style='normal')
    
    def _generate_speech(self, config: Dict[str, Any]):
        """生成发言稿"""
        # 标题
        title = config.get('title', '发言稿')
        self._add_paragraph(f"关于{title}的讲话", style='title')
        
        # 称呼
        self._add_paragraph("", style='normal')
        salutation = config.get('salutation', '尊敬的各位领导、同志们：')
        self._add_paragraph(salutation, style='normal')
        
        # 开场白
        introduction = config.get('introduction', '')
        if introduction:
            self._add_paragraph("", style='normal')
            self._add_paragraph(introduction, style='normal')
        
        # 正文
        sections = config.get('sections', [])
        for i, section in enumerate(sections, 1):
            self._add_paragraph("", style='normal')
            self._add_heading(f"一、{section.get('title', '')}", level=1)
            
            points = section.get('points', [])
            for j, point in enumerate(points, 1):
                self._add_heading(f"（一）{point.get('title', '')}", level=2)
                content = point.get('content', '')
                if content:
                    self._add_paragraph(content, style='normal')
        
        # 结束语
        conclusion = config.get('conclusion', '')
        if conclusion:
            self._add_paragraph("", style='normal')
            self._add_paragraph(conclusion, style='normal')
        
        # 结束语
        self._add_paragraph("", style='normal')
        self._add_paragraph("谢谢大家！", style='normal')
        
        # 落款
        self._add_paragraph("", style='normal')
        self._add_paragraph("", style='normal')
        organization = config.get('organization', '')
        if organization:
            self._add_paragraph(organization, style='normal')
        
        position = config.get('position', '')
        if position:
            self._add_paragraph(position, style='normal')
        
        name = config.get('name', '')
        if name:
            self._add_paragraph(name, style='normal')
        
        # 日期
        current_date = datetime.now().strftime("%Y年%m月%d日")
        self._add_paragraph(current_date, style='normal')
    
    def _generate_discussion_outline(self, config: Dict[str, Any]):
        """生成讨论提纲"""
        # 标题
        title = config.get('title', '讨论提纲')
        self._add_paragraph(f"{title}讨论提纲", style='title')
        
        # 讨论背景
        background = config.get('background', '')
        if background:
            self._add_heading("一、讨论背景", level=1)
            self._add_paragraph(background, style='normal')
        
        # 讨论目标
        objectives = config.get('objectives', [])
        if objectives:
            self._add_heading("二、讨论目标", level=1)
            for i, objective in enumerate(objectives, 1):
                self._add_paragraph(f"{i}. {objective}", style='normal')
        
        # 讨论要点
        points = config.get('points', [])
        if points:
            self._add_heading("三、讨论要点", level=1)
            for i, point in enumerate(points, 1):
                self._add_heading(f"（一）{point.get('title', '')}", level=2)
                subpoints = point.get('subpoints', [])
                for j, subpoint in enumerate(subpoints, 1):
                    self._add_paragraph(f"{j}. {subpoint}", style='normal')
        
        # 预期成果
        expected_outcomes = config.get('expected_outcomes', [])
        if expected_outcomes:
            self._add_heading("四、预期成果", level=1)
            for i, outcome in enumerate(expected_outcomes, 1):
                self._add_paragraph(f"{i}. {outcome}", style='normal')
        
        # 时间安排
        schedule = config.get('schedule', [])
        if schedule:
            self._add_heading("五、时间安排", level=1)
            for i, item in enumerate(schedule, 1):
                self._add_paragraph(f"{i}. {item}", style='normal')
        
        # 参考材料
        references = config.get('references', [])
        if references:
            self._add_heading("六、参考材料", level=1)
            for i, reference in enumerate(references, 1):
                self._add_paragraph(f"{i}. {reference}", style='normal')
    
    def _generate_work_report(self, config: Dict[str, Any]):
        """生成工作汇报"""
        # 标题
        title = config.get('title', '工作汇报')
        self._add_paragraph(f"关于{title}的工作汇报", style='title')
        
        # 主送机关
        recipient = config.get('recipient', '')
        if recipient:
            self._add_paragraph("", style='normal')
            self._add_paragraph(f"{recipient}：", style='normal')
        
        # 工作进展情况
        progress = config.get('progress', [])
        if progress:
            self._add_heading("一、工作进展情况", level=1)
            for i, item in enumerate(progress, 1):
                self._add_heading(f"（一）{item.get('area', '')}", level=2)
                details = item.get('details', [])
                for j, detail in enumerate(details, 1):
                    self._add_paragraph(f"{j}. {detail}", style='normal')
        
        # 主要成绩和亮点
        achievements = config.get('achievements', [])
        if achievements:
            self._add_heading("二、主要成绩和亮点", level=1)
            for i, achievement in enumerate(achievements, 1):
                self._add_paragraph(f"{i}. {achievement}", style='normal')
        
        # 存在的问题和困难
        problems = config.get('problems', [])
        if problems:
            self._add_heading("三、存在的问题和困难", level=1)
            for i, problem in enumerate(problems, 1):
                self._add_paragraph(f"{i}. {problem}", style='normal')
        
        # 下一步工作计划
        plans = config.get('plans', [])
        if plans:
            self._add_heading("四、下一步工作计划", level=1)
            for i, plan in enumerate(plans, 1):
                self._add_paragraph(f"{i}. {plan}", style='normal')
        
        # 需要支持的事项
        support_needed = config.get('support_needed', [])
        if support_needed:
            self._add_heading("五、需要支持的事项", level=1)
            for i, support in enumerate(support_needed, 1):
                self._add_paragraph(f"{i}. {support}", style='normal')
        
        # 结束语
        self._add_paragraph("", style='normal')
        self._add_paragraph("特此汇报。", style='normal')
        
        # 落款
        self._add_paragraph("", style='normal')
        self._add_paragraph("", style='normal')
        organization = config.get('organization', '')
        if organization:
            self._add_paragraph(organization, style='normal')
        
        # 日期
        current_date = datetime.now().strftime("%Y年%m月%d日")
        self._add_paragraph(current_date, style='normal')
    
    def _add_heading(self, text: str, level: int = 1):
        """添加标题"""
        if level == 1:
            self._add_paragraph(text, style='h1')
        elif level == 2:
            self._add_paragraph(text, style='h2')
    
    def _add_paragraph(self, text: str, style: str = 'normal'):
        """添加段落"""
        if not self.document:
            return
            
        p = self.document.add_paragraph(text)
        if style in self.styles:
            p.style = self.styles[style]
    
    def save_document(self, output_path: str):
        """保存文档"""
        if self.document:
            self.document.save(output_path)
            print(f"文档已保存到：{output_path}")
        else:
            print("错误：没有文档可保存")


def load_config(config_file: str) -> Dict[str, Any]:
    """加载配置文件"""
    if not os.path.exists(config_file):
        print(f"错误：配置文件不存在：{config_file}")
        sys.exit(1)
    
    try:
        with open(config_file, 'r', encoding='utf-8') as f:
            if config_file.endswith('.json'):
                return json.load(f)
            else:
                # 支持其他格式的配置文件
                print(f"错误：不支持的配置文件格式：{config_file}")
                sys.exit(1)
    except Exception as e:
        print(f"错误：加载配置文件失败：{e}")
        sys.exit(1)


def create_example_config(doc_type: str) -> Dict[str, Any]:
    """创建示例配置"""
    if doc_type == "meeting_minutes":
        return {
            "title": "数字化转型工作",
            "meeting_info": {
                "time": "2026年3月28日 14:00-16:00",
                "location": "第一会议室",
                "host": "张三",
                "recorder": "李四",
                "attendees": ["王五", "赵六", "孙七"],
                "absentees": ["周八"]
            },
            "topics": [
                "数字化转型工作方案讨论",
                "技术平台选型",
                "实施计划安排"
            ],
            "contents": [
                {
                    "topic": "数字化转型工作方案讨论",
                    "points": [
                        "分析了当前数字化发展形势",
                        "讨论了数字化转型的必要性",
                        "明确了数字化转型的目标"
                    ]
                }
            ],
            "resolutions": [
                {
                    "content": "原则通过《数字化转型工作方案》",
                    "responsible": "技术部",
                    "deadline": "2026年4月15日前"
                }
            ],
            "next_steps": [
                "技术部负责完善方案细节",
                "人力资源部负责人员培训安排",
                "财务部负责预算编制"
            ]
        }
    elif doc_type == "speech":
        return {
            "title": "推进数字化转型工作",
            "salutation": "尊敬的各位领导、同志们：",
            "introduction": "今天我们召开这次会议，主要目的是研究部署推进数字化转型工作。下面，我讲几点意见：",
            "sections": [
                {
                    "title": "充分认识数字化转型的重要意义",
                    "points": [
                        {
                            "title": "数字化转型是时代发展的必然要求",
                            "content": "当前，数字化浪潮席卷全球，新一代信息技术迅猛发展，数字化转型已成为企业发展的关键路径。"
                        }
                    ]
                }
            ],
            "conclusion": "同志们，数字化转型任务艰巨、责任重大。让我们以更加饱满的热情、更加务实的作风，扎实推进各项工作，确保数字化转型取得实效。",
            "organization": "某某单位",
            "position": "主任",
            "name": "张三"
        }
    else:
        return {}


def main():
    parser = argparse.ArgumentParser(description='官方文档生成器')
    parser.add_argument('--type', required=True, 
                       choices=['meeting_minutes', 'speech', 'discussion_outline', 'work_report'],
                       help='文档类型')
    parser.add_argument('--config', help='配置文件路径（JSON格式）')
    parser.add_argument('--output', required=True, help='输出文件路径')
    parser.add_argument('--example', action='store_true', help='使用示例配置')
    
    args = parser.parse_args()
    
    # 加载配置
    if args.example:
        config = create_example_config(args.type)
    elif args.config:
        config = load_config(args.config)
    else:
        print("错误：必须提供配置文件或使用--example参数")
        parser.print_help()
        sys.exit(1)
    
    # 生成文档
    try:
        generator = OfficialDocumentGenerator()
        generator.create_document(args.type, config)
        generator.save_document(args.output)
        print("文档生成成功！")
    except Exception as e:
        print(f"错误：生成文档失败：{e}")
        sys.exit(1)


if __name__ == "__main__":
    main()