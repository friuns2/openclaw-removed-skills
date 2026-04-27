#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
增强功能集成测试脚本
测试敏感词检查、修订历史、英文字体处理等所有增强功能
"""

import os
import sys
import json
import tempfile
import shutil
from datetime import datetime
from pathlib import Path

# 添加技能目录到路径
skill_dir = Path(__file__).parent.parent
sys.path.insert(0, str(skill_dir))


def test_sensitive_words_checker():
    """测试增强版敏感词检查器"""
    print("="*80)
    print("测试1：增强版敏感词检查器")
    print("="*80)
    
    try:
        from sensitive_words_check_enhanced import SensitiveWordsCheckerEnhanced
        
        # 创建测试数据目录
        test_data_dir = os.path.join(skill_dir, "test_data")
        os.makedirs(test_data_dir, exist_ok=True)
        
        # 创建检查器
        checker = SensitiveWordsCheckerEnhanced(test_data_dir)
        
        # 测试1：检查文本
        test_text = "这是一个测试文档，包含一些敏感词如台独和经济崩溃。"
        print(f"\n测试文本：{test_text}")
        
        result = checker.check_text(test_text, document_type='meeting_minutes')
        print(f"检查结果：{'通过' if result['valid'] else '未通过'}")
        print(f"敏感词数量：{result['sensitive_count']}")
        print(f"高级别敏感词：{result['high_count']}")
        print(f"中级别敏感词：{result['medium_count']}")
        
        # 测试2：检查不同文档类型
        print("\n测试不同文档类型的检查级别：")
        for doc_type, level in checker.document_type_levels.items():
            result = checker.check_text(test_text, document_type=doc_type)
            print(f"  {doc_type}: 级别={level}, 通过={result['valid']}")
        
        # 测试3：强制更新（模拟）
        print("\n测试强制更新敏感词库：")
        success = checker.force_update()
        print(f"强制更新结果：{'成功' if success else '失败'}")
        
        # 测试4：导出词库
        export_file = os.path.join(test_data_dir, "export_test.json")
        success = checker.export_words(export_file)
        print(f"导出词库：{'成功' if success else '失败'}")
        
        if success and os.path.exists(export_file):
            print(f"导出文件大小：{os.path.getsize(export_file)} 字节")
        
        return True
        
    except Exception as e:
        print(f"敏感词检查器测试失败：{e}")
        import traceback
        traceback.print_exc()
        return False


def test_revision_history():
    """测试修订历史管理器"""
    print("\n" + "="*80)
    print("测试2：修订历史管理器")
    print("="*80)
    
    try:
        from revision_history import RevisionHistoryManager, DocumentRevisionTracker
        
        # 创建测试目录
        test_history_dir = os.path.join(skill_dir, "test_history")
        os.makedirs(test_history_dir, exist_ok=True)
        
        # 创建管理器
        manager = RevisionHistoryManager(test_history_dir)
        tracker = DocumentRevisionTracker(manager)
        
        # 测试文档路径
        test_doc_path = os.path.join(test_history_dir, "test_document.docx")
        
        # 测试1：创建文档记录
        print("\n测试文档创建记录：")
        create_record = tracker.track_document_creation(
            document_path=test_doc_path,
            document_type='meeting_minutes',
            author='测试用户',
            metadata={'purpose': '功能测试'}
        )
        
        print(f"创建记录ID：{create_record['id']}")
        print(f"文档类型：{create_record['document_type']}")
        print(f"操作类型：{create_record['action']}")
        print(f"作者：{create_record['author']}")
        
        # 测试2：修改文档记录
        print("\n测试文档修改记录：")
        changes = [
            {'type': 'content_update', 'description': '更新了会议议题'},
            {'type': 'format_adjust', 'description': '调整了段落格式'}
        ]
        
        modify_record = tracker.track_document_modification(
            document_path=test_doc_path,
            document_type='meeting_minutes',
            author='测试用户',
            changes=changes,
            metadata={'reason': '内容优化'}
        )
        
        print(f"修改记录ID：{modify_record['id']}")
        print(f"变更数量：{len(changes)}")
        
        # 测试3：保存所有历史
        print("\n测试保存所有历史：")
        saved_files = tracker.save_all_histories(output_format='json')
        print(f"保存的历史文件数量：{len(saved_files)}")
        for i, file in enumerate(saved_files, 1):
            print(f"  {i}. {file}")
        
        # 测试4：获取文档摘要
        print("\n测试文档修订摘要：")
        summary = tracker.get_document_summary(test_doc_path)
        print(f"修订次数：{summary.get('revision_count', 0)}")
        print(f"首次修订：{summary.get('first_revision', '无')}")
        print(f"最后修订：{summary.get('latest_revision', '无')}")
        
        # 测试5：不同格式输出
        print("\n测试不同格式输出：")
        test_outputs = []
        for format_type in ['json', 'txt', 'md', 'csv']:
            output_file = manager.save_revision_history(
                document_path=test_doc_path,
                records=[create_record, modify_record],
                output_format=format_type,
                include_all=True
            )
            test_outputs.append(output_file)
            print(f"  {format_type.upper()}格式：{output_file}")
        
        # 清理测试文件
        for file in test_outputs:
            if os.path.exists(file):
                os.remove(file)
        
        # 清理测试目录
        if os.path.exists(test_history_dir):
            shutil.rmtree(test_history_dir)
        
        return True
        
    except Exception as e:
        print(f"修订历史管理器测试失败：{e}")
        import traceback
        traceback.print_exc()
        return False


def test_enhanced_document_generator():
    """测试增强版文档生成器"""
    print("\n" + "="*80)
    print("测试3：增强版文档生成器")
    print("="*80)
    
    try:
        from generate_document_enhanced import EnhancedDocumentGenerator
        
        # 创建测试目录
        test_output_dir = os.path.join(skill_dir, "test_output")
        os.makedirs(test_output_dir, exist_ok=True)
        
        # 测试配置
        test_config = {
            "title": "数字化转型工作",
            "meeting_info": {
                "time": "2026年3月28日 14:00-16:00",
                "location": "第一会议室",
                "host": "张三",
                "recorder": "李四",
                "attendees": ["王五", "赵六", "孙七"],
                "absentees": ["周八"]
            },
            "topics": [
                "数字化转型工作方案讨论",
                "技术平台选型 (Technology Platform Selection)",
                "实施计划安排 (Implementation Plan)"
            ],
            "contents": [
                {
                    "topic": "数字化转型工作方案讨论",
                    "points": [
                        "分析了当前数字化发展形势",
                        "讨论了数字化转型的必要性 (The necessity of digital transformation)",
                        "明确了数字化转型的目标 (Digital transformation goals)"
                    ]
                }
            ],
            "resolutions": [
                {
                    "content": "原则通过《数字化转型工作方案》",
                    "responsible": "技术部 (Technology Department)",
                    "deadline": "2026年4月15日前"
                }
            ],
            "next_steps": [
                "技术部负责完善方案细节",
                "人力资源部负责人员培训安排 (HR Department arranges training)",
                "财务部负责预算编制 (Budget preparation)"
            ]
        }
        
        # 测试1：创建文档（启用历史跟踪）
        print("\n测试文档创建（启用历史跟踪）：")
        generator = EnhancedDocumentGenerator(enable_history_tracking=True)
        
        output_path = os.path.join(test_output_dir, "test_meeting_minutes.docx")
        generator.create_document(
            doc_type="meeting_minutes",
            config=test_config,
            author="测试用户",
            metadata={"test": True, "version": "2.0"}
        )
        
        # 保存文档
        success = generator.save_document(output_path, "测试用户")
        print(f"文档生成：{'成功' if success else '失败'}")
        
        if success and os.path.exists(output_path):
            file_size = os.path.getsize(output_path)
            print(f"文档大小：{file_size} 字节")
            
            # 检查文档内容
            import docx
            doc = docx.Document(output_path)
            print(f"文档段落数：{len(doc.paragraphs)}")
            
            # 检查中英文内容
            english_count = 0
            for para in doc.paragraphs:
                if re.search(r'[a-zA-Z]', para.text):
                    english_count += 1
            
            print(f"包含英文的段落数：{english_count}")
        
        # 测试2：获取修订摘要
        print("\n测试修订摘要：")
        summary = generator.get_revision_summary(output_path)
        if summary:
            print(f"修订次数：{summary.get('revision_count', 0)}")
            print(f"最后操作：{summary.get('latest_action', '无')}")
        else:
            print("未找到修订摘要")
        
        # 测试3：创建文档（禁用历史跟踪）
        print("\n测试文档创建（禁用历史跟踪）：")
        generator_no_history = EnhancedDocumentGenerator(enable_history_tracking=False)
        
        output_path2 = os.path.join(test_output_dir, "test_speech.docx")
        generator_no_history.create_document(
            doc_type="speech",
            config=test_config,
            author="测试用户"
        )
        
        success2 = generator_no_history.save_document(output_path2, "测试用户")
        print(f"文档生成（无历史）：{'成功' if success2 else '失败'}")
        
        # 清理测试文件
        if os.path.exists(test_output_dir):
            for file in os.listdir(test_output_dir):
                if file.endswith('.docx'):
                    os.remove(os.path.join(test_output_dir, file))
            os.rmdir(test_output_dir)
        
        return True
        
    except Exception as e:
        print(f"增强版文档生成器测试失败：{e}")
        import traceback
        traceback.print_exc()
        return False


def test_wps_compatibility():
    """测试WPS兼容性"""
    print("\n" + "="*80)
    print("测试4：WPS兼容性检查")
    print("="*80)
    
    try:
        from generate_document_enhanced import EnhancedDocumentGenerator
        
        # 创建测试配置（包含中英文混合内容）
        test_config = {
            "title": "WPS兼容性测试",
            "meeting_info": {
                "time": "2026年3月28日",
                "location": "测试会议室",
                "host": "测试主持人",
                "recorder": "测试记录员"
            },
            "topics": [
                "中英文混合测试 (Chinese-English Mixed Test)",
                "字体兼容性测试",
                "格式保持测试"
            ]
        }
        
        # 创建文档
        generator = EnhancedDocumentGenerator(enable_history_tracking=False)
        generator.create_document("meeting_minutes", test_config, "WPS测试")
        
        # 保存文档
        test_file = "wps_compatibility_test.docx"
        generator.save_document(test_file, "WPS测试")
        
        if os.path.exists(test_file):
            print(f"✅ 测试文档已创建：{test_file}")
            print(f"📊 文件大小：{os.path.getsize(test_file)} 字节")
            
            # 检查文档属性
            import docx
            doc = docx.Document(test_file)
            
            print("\n📋 文档属性检查：")
            print(f"   段落数量：{len(doc.paragraphs)}")
            print(f"   章节数量：{len(doc.sections)}")
            
            # 检查字体设置
            print("\n🔤 字体设置检查：")
            styles = doc.styles
            for style_name in ['Normal', 'Title_Style', 'Heading1_Style']:
                if style_name in styles:
                    style = styles[style_name]
                    print(f"   {style_name}: {style.font.name}")
            
            # 检查页面设置
            print("\n📄 页面设置检查：")
            section = doc.sections[0]
            print(f"   页边距：上{section.top_margin.cm:.1f}cm, 下{section.bottom_margin.cm:.1f}cm")
            print(f"           左{section.left_margin.cm:.1f}cm, 右{section.right_margin.cm:.1f}cm")
            print(f"   纸张大小：{section.page_width.cm:.1f}cm × {section.page_height.cm:.1f}cm")
            
            # 清理测试文件
            os.remove(test_file)
            print(f"\n🧹 已清理测试文件：{test_file}")
        
        print("\n💡 WPS兼容性建议：")
        print("   1. 确保使用标准字体（宋体、黑体、楷体、仿宋）")
        print("   2. 英文字体统一使用 Times New Roman")
        print("   3. 避免使用特殊字符和复杂格式")
        print("   4. 在WPS中打开后检查格式是否正常")
        
        return True
        
    except Exception as e:
        print(f"WPS兼容性测试失败：{e}")
        return False


def test_integration():
    """集成测试所有功能"""
    print("\n" + "="*80)
    print("集成测试：所有增强功能")
    print("="*80)
    
    # 创建临时测试目录
    temp_dir = tempfile.mkdtemp(prefix="official_doc_test_")
    print(f"临时测试目录：{temp_dir}")
    
    test_results = {
        'sensitive_words_checker': False,
        'revision_history': False,
        'enhanced_document_generator': False,
        'wps_compatibility': False
    }
    
    try:
        # 运行所有测试
        test_results['sensitive_words_checker'] = test_sensitive_words_checker()
        test_results['revision_history'] = test_revision_history()
        test_results['enhanced_document_generator'] = test_enhanced_document_generator()
        test_results['wps_compatibility'] = test_wps_compatibility()
        
        # 显示测试结果
        print("\n" + "="*80)
        print("测试结果汇总")
        print("="*80)
        
        total_tests = len(test_results)
        passed_tests = sum(1 for result in test_results.values() if result)
        
        for test_name, result in test_results.items():
            status = "✅ 通过" if result else "❌ 失败"
            print(f"{test_name}: {status}")
        
        print(f"\n📊 总计：{passed_tests}/{total_tests} 项测试通过")
        
        if passed_tests == total_tests:
            print("\n🎉 所有测试通过！增强功能正常。")
        else:
            print(f"\n⚠️  {total_tests - passed_tests} 项测试失败，请检查相关问题。")
        
    finally:
        # 清理临时目录
        if os.path.exists(temp_dir):
            shutil.rmtree(temp_dir)
            print(f"\n🧹 已清理临时目录：{temp_dir}")
    
    return all(test_results.values())


def main():
    """主函数"""
    print("官方文档生成器 - 增强功能测试套件")
    print("="*80)
    print("测试内容包括：")
    print("  1. 增强版敏感词检查器（联网更新、文档类型自适应）")
    print("  2. 修订历史管理器（JSON/TXT格式记录）")
    print("  3. 增强版文档生成器（英文字体处理）")
    print("  4. WPS兼容性检查")
    print("="*80)
    
    try:
        # 检查依赖
        print("\n🔍 检查依赖...")
        try:
            import docx
            print("✅ python-docx 已安装")
        except ImportError:
            print("❌ python-docx 未安装")
            print("   请运行：pip install python-docx")
            return False
        
        # 运行集成测试
        success = test_integration()
        
        if success:
            print("\n" + "="*80)
            print("🏆 所有测试完成！增强功能准备就绪。")
            print("="*80)
            print("\n📋 可用功能：")
            print("  • 敏感词检查（支持联网更新、文档类型自适应）")
            print("  • 修订历史记录（独立的JSON/TXT文件）")
            print("  • 中英文混合处理（统一使用Times New Roman）")
            print("  • WPS兼容性优化")
            print("  • 多种文档类型支持")
        else:
            print("\n" + "="*80)
            print("⚠️ 测试失败，部分功能可能不可用。")
            print("="*80)
        
        return success
        
    except Exception as e:
        print(f"\n❌ 测试过程中发生错误：{e}")
        import traceback
        traceback.print_exc()
        return False


# 导入re模块用于测试
import re

if __name__ == "__main__":
    success = main()
    sys.exit(0 if success else 1)