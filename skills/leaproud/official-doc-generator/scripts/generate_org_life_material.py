#!/usr/bin/env python3
# -*- coding: utf-8 -*-
'''
组织生活会个人对照检查材料生成器
生成符合公文格式（GB/T 9704-2012）的Word文档
'''

import os
import sys
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    print('错误：需要安装 python-docx 库')
    sys.exit(1)


# ============================================================
# 文档内容
# ============================================================

TITLE = '组织生活会个人对照检查材料'

SECTION_1_TITLE = '一、对照检查剖析'

POLITICAL_LOYALTY = (
    '在政治忠诚方面，本人始终坚持以习近平新时代中国特色社会主义思想为指导，'
    '认真学习贯彻党的路线方针政策，坚决维护党中央权威和集中统一领导。'
    '但对照更高标准，还存在以下不足：一是理论学习的系统性不够，有时满足于完成规定学习任务，'
    '对部分新理论、新政策的钻研深度不够，存在\u201c泛泛而学\u201d的现象；'
    '二是理论联系实际不够紧密，将政治理论学习成果转化为推动工作、解决实际问题的能力还有差距，'
    '有时存在学用脱节的问题；三是在重大原则问题上虽能站稳立场，但主动发声亮剑的意识还不够强，'
    '对身边一些错误言论和模糊认识有时未能及时予以纠正。'
)

PARTY_SPIRIT = (
    '在党性锤炼方面，本人能够自觉参加组织生活，认真落实\u201c三会一课\u201d制度，'
    '积极开展批评与自我批评。但深刻反思，还存在以下差距：一是党性锻炼的主动性有所欠缺，'
    '有时存在\u201c过得去\u201d的思想，在自我净化、自我完善、自我革新、自我提高方面做得还不够；'
    '二是宗旨意识树得不够牢固，在日常工作中有时未能充分站在服务对象的角度思考问题，'
    '特别是面对繁杂事务时，容易产生急躁情绪，服务意识有待进一步增强；'
    '三是艰苦奋斗精神有所弱化，面对工作中的困难和压力时，有时存在畏难情绪，'
    '攻坚克难的劲头还不够足。'
)

DISCIPLINE_RULES = (
    '在纪律规矩方面，本人严格遵守党的政治纪律、组织纪律、廉洁纪律、群众纪律、'
    '工作纪律和生活纪律，自觉执行中央八项规定及其实施细则精神。'
    '但认真对照检查，还存在以下问题：一是纪律意识的自觉性还不够高，'
    '有时认为只要不触碰底线就行，对自身要求有所降低，对一些\u201c小节\u201d问题不够重视；'
    '二是执行规章制度有时存在打折扣现象，个别工作流程上存在简化环节的情况，'
    '未能做到百分之百严格执行；三是请示报告制度落实不够到位，'
    '个别事项存在事后报告或漏报的情况，政治敏锐性需要进一步提高。'
)

IMPROVE_STYLE = (
    '在作风建设方面，本人注重加强自身作风修养，努力做到勤勉务实、廉洁自律。'
    '但对照新时代作风建设要求，还存在以下不足：一是工作作风还不够扎实，'
    '有时存在形式主义倾向，重留痕轻实效的问题尚未完全根治，'
    '个别工作满足于\u201c做了\u201d而非\u201c做好了\u201d；'
    '二是调查研究不够深入，下基层了解实际情况的频次和深度都有所不足，'
    '掌握第一手资料不够全面，导致部分决策缺乏充分依据；三是担当精神还不够强，'
    '面对矛盾和困难时，有时存在绕着走、等靠要的思想，缺乏主动攻坚、迎难而上的勇气和魄力。'
)

LOGISTICS_SECTION_TITLE = '（五）后勤领域专项整肃治理'

LOGISTICS_1_LOYALTY = (
    '政治忠诚方面：能够坚决贯彻上级关于后勤工作的各项决策部署，在大是大非面前保持清醒头脑。'
    '但对照专项整肃要求，一是对后勤领域政治属性的认识还不够深刻，'
    '有时将后勤工作简单等同于事务性工作，对其政治意义和战略地位理解不够到位；'
    '二是在后勤保障工作中贯彻政治要求的意识需要加强，'
    '特别是在物资采购、工程建设等敏感领域，政治把关的能力还有待提高。'
)

LOGISTICS_2_PRINCIPLE = (
    '党性原则方面：能够坚持党性原则，自觉做到公私分明。但深刻反思，'
    '一是在后勤资源分配中，有时存在平均主义思想，未能充分体现奖优罚劣的原则导向；'
    '二是面对后勤工作中的人际关系压力时，坚持原则的定力还不够，'
    '在个别事项处理上存在\u201c和稀泥\u201d的现象。'
)

LOGISTICS_3_POWER = (
    '履职用权方面：能够按照规定权限行使后勤管理职责。但对照检查，'
    '一是在后勤审批流程中，有时存在效率意识过强、程序意识不足的问题，'
    '个别环节的把关不够严格；二是后勤信息公开透明度有待提高，'
    '部分经费使用情况的公示还不够及时、不够详细。'
)

LOGISTICS_4_DISCIPLINE = (
    '遵规守纪方面：能够遵守后勤领域各项规章制度。但还存在以下不足：'
    '一是对后勤领域廉政风险点的排查还不够全面，风险防控措施还有漏洞；'
    '二是在物资管理方面，出入库登记有时不够及时，库存盘点存在走过场的情况，'
    '精细化管理水平需要提升。'
)

LOGISTICS_5_ETHICS = (
    '行业操守方面：注重保持后勤工作者的职业操守。但对照要求，'
    '一是在与供应商接触中，边界意识还需进一步强化，'
    '个别场合的交往活动超出了正常工作范畴；'
    '二是后勤服务标准意识不够强，存在\u201c差不多就行\u201d的心态，'
    '精细化和品质化服务水平有待提升。'
)

LOGISTICS_6_TRADITION = (
    '赓续传统方面：能够继承和发扬勤俭节约、艰苦奋斗的优良传统。但深刻反思，'
    '一是在新形势下对后勤工作传统的传承创新不够，'
    '存在\u201c老办法管用、新办法不会用\u201d的问题；'
    '二是厉行节约的意识有所淡化，在日常办公耗材使用、水电管理等细节方面存在浪费现象，'
    '过紧日子的思想树立得还不够牢固。'
)

SECTION_2_TITLE = '二、下一步整改措施'

MEASURES = [
    '强化理论武装，筑牢思想根基。坚持把政治理论学习放在首位，制定个人年度学习计划，'
    '每周至少安排两次集中学习时间，系统研读党的创新理论原著原文，力求学深悟透、融会贯通。'
    '注重理论联系实际，每季度撰写一篇学习心得，切实把学习成果转化为推动工作的实际能力。'
    '在重大原则问题上敢于亮剑、善于发声，坚决同各种错误思想作斗争。',

    '锤炼党性修养，践行根本宗旨。自觉加强党性锻炼，主动参加党内政治生活，'
    '认真开展批评与自我批评，做到红脸出汗、排毒治病。牢固树立以服务对象为中心的工作理念，'
    '经常深入基层一线了解实际需求，每季度至少开展一次专项调研，'
    '切实解决服务对象反映强烈的突出问题。传承和弘扬艰苦奋斗精神，'
    '面对困难不退缩，面对压力不低头，始终保持昂扬向上的精神状态。',

    '严守纪律底线，筑牢廉洁防线。持续加强对党规党纪的学习，做到知敬畏、存戒惧、守底线。'
    '严格执行请示报告制度，做到事前请示、事后报告，杜绝任何漏报、迟报现象。'
    '完善后勤领域廉政风险防控机制，全面排查风险点，建立台账、逐项整改。'
    '自觉净化社交圈、生活圈、朋友圈，保持清正廉洁的政治本色。',

    '改进工作作风，提升服务效能。坚决整治形式主义问题，把精力集中到解决实际问题上来，'
    '不搞花架子、不做表面文章。大兴调查研究之风，坚持深入一线、深入实际，'
    '每年不少于两次专项调研，确保决策有依据、工作有实效。'
    '强化担当意识，面对矛盾问题敢于碰硬、善于攻坚，做到守土有责、守土尽责。',

    '深化后勤整肃，规范权力运行。严格落实后勤领域专项整肃治理各项要求，'
    '对照\u201c六个方面\u201d逐项整改销号。完善后勤管理制度体系，'
    '规范采购、工程、物资管理等关键环节的操作流程，做到有章可循、有据可查。'
    '推进后勤信息公开，定期公示经费使用和物资管理情况，接受群众监督。'
    '坚持勤俭办一切事业，严格执行节约制度，杜绝铺张浪费现象。'
]


class OrgLifeDocGenerator:
    '''组织生活会材料文档生成器'''

    def __init__(self):
        self.doc = Document()
        self._setup_page()
        self._setup_styles()

    def _setup_page(self):
        sec = self.doc.sections[0]
        sec.page_width = Cm(21.0)
        sec.page_height = Cm(29.7)
        # GB/T 9704-2012: 天头37mm, 订口28mm, 版心156mm x 225mm
        # 下白边 = 297 - 37 - 225 = 35mm, 右白边 = 210 - 28 - 156 = 26mm
        sec.top_margin = Cm(3.7)
        sec.bottom_margin = Cm(3.5)
        sec.left_margin = Cm(2.8)
        sec.right_margin = Cm(2.6)

    def _setup_styles(self):
        # GB/T 9704-2012 5.2.2: 如无特殊说明,公文格式各要素一般用3号仿宋体字
        # 3号 = 16pt
        # 5.2.3: 一般每面排22行,每行排28个字
        # "一行" = 一个汉字高度(16pt) + 3号汉字高度的7/8(14pt) ≈ 30pt
        from docx.enum.text import WD_LINE_SPACING
        style_normal = self.doc.styles['Normal']
        style_normal.font.name = '\u4eff\u5b8b'
        style_normal.font.size = Pt(16)
        # 国标行距：固定值28.95pt (16 + 16*7/8 ≈ 30pt，取28.95pt保证每面22行)
        style_normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        style_normal.paragraph_format.line_spacing = Pt(28.95)
        # 7.3.3: 每个自然段左空二字
        style_normal.paragraph_format.first_line_indent = Cm(1.12)  # 2 x 3号字宽(≈5.6mm)
        rpr = style_normal._element.get_or_add_rPr()
        rpr.rFonts.set(qn('w:eastAsia'), '\u4eff\u5b8b')
        rpr.rFonts.set(qn('w:ascii'), 'Times New Roman')
        rpr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')

    def _set_run_font(self, run, east_asia, size_pt):
        '''设置run的中英文字体和字号'''
        run.font.name = east_asia
        run.font.size = Pt(size_pt)
        rpr = run._r.get_or_add_rPr()
        rpr.rFonts.set(qn('w:eastAsia'), east_asia)
        rpr.rFonts.set(qn('w:ascii'), 'Times New Roman')
        rpr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')

    def _add_title(self, text):
        # 7.3.1: 一般用2号小标宋体字,居中排布
        from docx.enum.text import WD_LINE_SPACING
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(28.95)
        run = p.add_run(text)
        self._set_run_font(run, '\u65b9\u6b63\u5c0f\u6807\u5b8b\u7b80\u4f53', 22)
        return p

    def _add_author(self, author_text):
        '''署名：标题下方，3号楷体_GB2312，居中，行距不变'''
        from docx.enum.text import WD_LINE_SPACING
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(28.95)
        p.paragraph_format.first_line_indent = Cm(0)
        run = p.add_run(author_text)
        self._set_run_font(run, '\u6977\u4f53_GB2312', 16)
        return p

    def _add_section_heading(self, text):
        '''一级标题：7.3.3 第一层用黑体字，不缩进'''
        from docx.enum.text import WD_LINE_SPACING
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(28.95)
        p.paragraph_format.first_line_indent = Cm(0)
        run = p.add_run(text)
        self._set_run_font(run, '\u9ed1\u4f53', 16)
        return p

    def _add_sub_heading(self, text):
        '''二级标题：7.3.3 第二层用楷体字，缩进2字'''
        from docx.enum.text import WD_LINE_SPACING
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(28.95)
        p.paragraph_format.first_line_indent = Cm(1.12)
        run = p.add_run(text)
        self._set_run_font(run, '\u6977\u4f53', 16)
        return p

    def _add_body(self, text):
        '''正文段落：3号仿宋_GB2312，首行缩进2字，固定行距28.95pt'''
        from docx.enum.text import WD_LINE_SPACING
        p = self.doc.add_paragraph()
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(28.95)
        p.paragraph_format.first_line_indent = Cm(1.12)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text)
        self._set_run_font(run, '\u4eff\u5b8b_GB2312', 16)
        return p

    def build(self, title=None, author=None):
        doc_title = title or TITLE
        self._add_title(doc_title)
        if author:
            self._add_author(author)

        # 第一部分
        self._add_section_heading(SECTION_1_TITLE)

        self._add_sub_heading('（一）筑牢政治忠诚')
        self._add_body(POLITICAL_LOYALTY)

        self._add_sub_heading('（二）加强党性锤炼')
        self._add_body(PARTY_SPIRIT)

        self._add_sub_heading('（三）严守纪律规矩')
        self._add_body(DISCIPLINE_RULES)

        self._add_sub_heading('（四）改作风树新风')
        self._add_body(IMPROVE_STYLE)

        self._add_sub_heading(LOGISTICS_SECTION_TITLE)

        self._add_body('1.政治忠诚')
        self._add_body(LOGISTICS_1_LOYALTY)

        self._add_body('2.党性原则')
        self._add_body(LOGISTICS_2_PRINCIPLE)

        self._add_body('3.履职用权')
        self._add_body(LOGISTICS_3_POWER)

        self._add_body('4.遵规守纪')
        self._add_body(LOGISTICS_4_DISCIPLINE)

        self._add_body('5.行业操守')
        self._add_body(LOGISTICS_5_ETHICS)

        self._add_body('6.赓续传统')
        self._add_body(LOGISTICS_6_TRADITION)

        # 第二部分
        self._add_section_heading(SECTION_2_TITLE)

        for i, measure in enumerate(MEASURES, 1):
            self._add_body(f'{i}.{measure}')

    def save(self, output_path):
        self.doc.save(output_path)
        print(f'文档已保存: {output_path}')
        total = sum(len(p.text) for p in self.doc.paragraphs)
        print(f'总字数: 约 {total} 字')


def main():
    import argparse
    parser = argparse.ArgumentParser(description='组织生活会个人对照检查材料生成器')
    parser.add_argument('--output-dir', required=True, help='输出目录（必须指定，不允许默认路径）')
    parser.add_argument('--filename', default=None, help='输出文件名（不含路径）')
    parser.add_argument('--title', default=None, help='文档标题')
    parser.add_argument('--author', default=None, help='署名（如"XXXX单位 张三"）')
    args = parser.parse_args()

    output_dir = args.output_dir
    os.makedirs(output_dir, exist_ok=True)

    date_str = datetime.now().strftime('%Y%m%d')
    if args.filename:
        filename = args.filename
    else:
        filename = f'\u7ec4\u7ec7\u751f\u6d3b\u4f1a\u4e2a\u4eba\u5bf9\u7167\u68c0\u67e5\u6750\u6599_{date_str}.docx'
    output_path = os.path.join(output_dir, filename)

    gen = OrgLifeDocGenerator()
    gen.build(title=args.title, author=args.author)
    gen.save(output_path)
    print(f'\u5b8c\u6210\uff01\u6587\u4ef6\u4f4d\u7f6e: {output_path}')


if __name__ == '__main__':
    main()
