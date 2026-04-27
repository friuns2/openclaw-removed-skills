#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
增强版官方文档生成器
支持英文字体统一处理（Times New Roman）和修订历史跟踪
"""

import argparse
import json
import os
import sys
import re
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, Optional, List, Tuple

# 添加技能目录到路径
skill_dir = Path(__file__).parent.parent
sys.path.insert(0, str(skill_dir))

try:
    from docx import Document
    from docx.shared import Pt, Cm, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    print("错误：需要安装 python-docx 库")
    print("请运行：pip install python-docx")
    sys.exit(1)

# 导入修订历史管理器
try:
    from revision_history import RevisionHistoryManager, DocumentRevisionTracker
except ImportError:
    print("警告：无法导入修订历史模块，修订历史功能将不可用")
    RevisionHistoryManager = None
    DocumentRevisionTracker = None


class EnhancedDocumentGenerator:
    """增强版官方文档生成器"""
    
    def __init__(self, enable_history_tracking: bool = True):
        """
        初始化文档生成器
        
        Args:
            enable_history_tracking: 是否启用修订历史跟踪
        """
        self.document = None
        self.styles = {}
        self.enable_history_tracking = enable_history_tracking
        
        # 初始化修订历史跟踪器
        self.history_manager = None
        self.revision_tracker = None
        if enable_history_tracking and RevisionHistoryManager:
            self.history_manager = RevisionHistoryManager()
            self.revision_tracker = DocumentRevisionTracker(self.history_manager)
        
        # 英文字体设置
        self.english_font_name = "Times New Roman"
        
        # 标题字体名称（统一使用方正小标宋简体）
        self.title_font_name = '方正小标宋简体'
        
        # 中文字体映射（主字体 -> 备用字体）
        self.chinese_font_mapping = {
            '小标宋': ['方正小标宋简体', '小标宋', 'SimSun', '宋体'],
            '黑体': ['黑体', 'SimHei', '微软雅黑'],
            '楷体': ['楷体_GB2312', '楷体', 'KaiTi'],
            '仿宋': ['仿宋_GB2312', '仿宋', 'FangSong']
        }
    
    def create_document(self, 
                       doc_type: str, 
                       config: Dict[str, Any],
                       author: str = "系统生成",
                       metadata: Optional[Dict] = None) -> Document:
        """创建文档
        
        Args:
            doc_type: 文档类型
            config: 配置参数
            author: 作者
            metadata: 额外元数据
            
        Returns:
            Document: Word文档对象
        """
        # 创建新文档
        self.document = Document()
        
        # 设置页面格式
        self._setup_page_format()
        
        # 添加标准页码
        for section in self.document.sections:
            self._add_page_numbers(section)
        
        # 在 settings.xml 中强制添加 evenAndOddHeaders（WPS兼容）
        settings_el = self.document.settings.element
        existing = settings_el.find(qn('w:evenAndOddHeaders'))
        if existing is None:
            even_odd = OxmlElement('w:evenAndOddHeaders')
            settings_el.append(even_odd)
        
        # 设置样式（包括英文字体）
        self._setup_enhanced_styles()
        
        # 生成文档内容
        self._generate_document_content(doc_type, config)
        
        # 记录文档创建
        if self.enable_history_tracking and self.revision_tracker:
            doc_metadata = metadata or {}
            doc_metadata.update({
                'document_type': doc_type,
                'config_summary': self._summarize_config(config),
                'generator_version': '2.0'
            })
            
            # 暂时保存文档以计算哈希（使用系统临时目录，避免写入 skill 目录）
            import tempfile
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tf:
                temp_path = tf.name
            self.document.save(temp_path)
            
            # 创建修订记录
            revision_record = self.revision_tracker.track_document_creation(
                document_path=temp_path,
                document_type=doc_type,
                author=author,
                metadata=doc_metadata
            )
            
            print(f"📝 文档创建记录：{revision_record['id']}")
            
            # 清理临时文件
            if os.path.exists(temp_path):
                os.remove(temp_path)
        
        return self.document
    
    def _setup_page_format(self):
        """设置页面格式（GB/T 9704-2012标准）"""
        from docx.enum.text import WD_LINE_SPACING
        sections = self.document.sections
        for section in sections:
            # GB/T 9704-2012 5.2.1: 天头37mm, 订口28mm, 版心156mm x 225mm
            section.top_margin = Cm(3.7)
            section.bottom_margin = Cm(3.5)
            section.left_margin = Cm(2.8)
            section.right_margin = Cm(2.6)

            # A4纸
            section.page_width = Cm(21.0)
            section.page_height = Cm(29.7)

            # 页码：单页居右空一字，双页居左空一字
            section.different_first_page_header_footer = False
            section.odd_and_even_pages_header_footer = True

            # 用原始XML强制确保 oddAndEven 标记写入（WPS兼容）
            sectPr = section._sectPr
            if sectPr is not None:
                pgSz = sectPr.find(qn('w:pgSz'))
                if pgSz is not None:
                    # 确保pgSz存在后再获取父节点
                    parent = sectPr
                else:
                    parent = sectPr
                # 查找或创建 titlePg 元素（控制首页不同）
                # 查找或创建 evenAndOddHeaders 元素（控制奇偶页不同）
                existing = sectPr.find(qn('w:evenAndOddHeaders'))
                if existing is None:
                    even_odd = OxmlElement('w:evenAndOddHeaders')
                    sectPr.insert(0, even_odd)
    
    def _add_page_numbers(self, section):
        """添加标准页码（GB/T 9704-2012: 4号半角宋体，一字线，版心下7mm，单页右/双页左）"""
        try:
            def _setup_page_number(paragraph, alignment):
                """在一个页脚段落中设置一字线 + PAGE 域"""
                paragraph.text = ""
                paragraph.alignment = alignment
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                paragraph.paragraph_format.line_spacing = Pt(14)

                r1 = paragraph.add_run("— ")
                r1.font.name = '宋体'
                r1.font.size = Pt(14)
                r1._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

                rn = paragraph.add_run()
                rn.font.name = '宋体'
                rn.font.size = Pt(14)
                rn._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                
                fldBegin = OxmlElement('w:fldChar')
                fldBegin.set(qn('w:fldCharType'), 'begin')
                instr = OxmlElement('w:instrText')
                instr.set(qn('xml:space'), 'preserve')
                instr.text = " PAGE "
                fldEnd = OxmlElement('w:fldChar')
                fldEnd.set(qn('w:fldCharType'), 'end')
                rn._r.append(fldBegin)
                rn._r.append(instr)
                rn._r.append(fldEnd)

                r2 = paragraph.add_run(" —")
                r2.font.name = '宋体'
                r2.font.size = Pt(14)
                r2._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

            # 奇数页页脚（居右空一字）
            odd_footer = section.footer
            odd_footer.is_linked_to_previous = False
            _setup_page_number(odd_footer.paragraphs[0], WD_ALIGN_PARAGRAPH.RIGHT)

            # 偶数页页脚（居左空一字）
            even_footer = section.even_page_footer
            even_footer.is_linked_to_previous = False
            _setup_page_number(even_footer.paragraphs[0], WD_ALIGN_PARAGRAPH.LEFT)

        except Exception as e:
            print(f"添加页码失败：{e}")
    
    def _add_header(self, section):
        """添加页眉"""
        try:
            header = section.header
            header_paragraph = header.paragraphs[0]
            header_paragraph.text = ""
            
            # 可以在这里添加单位名称等
            # header_paragraph.text = "某某单位"
            # header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        except Exception as e:
            print(f"添加页眉失败：{e}")
    
    def _add_footer(self, section):
        """添加页脚（页码）"""
        try:
            footer = section.footer
            footer_paragraph = footer.paragraphs[0]
            footer_paragraph.text = ""
            footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加页码
            run = footer_paragraph.add_run()
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            
            instrText = OxmlElement('w:instrText')
            instrText.text = "PAGE"
            
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')
            
            run._r.append(fldChar1)
            run._r.append(instrText)
            run._r.append(fldChar2)
            
            run = footer_paragraph.add_run(" / ")
            
            run = footer_paragraph.add_run()
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            
            instrText = OxmlElement('w:instrText')
            instrText.text = "NUMPAGES"
            
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')
            
            run._r.append(fldChar1)
            run._r.append(instrText)
            run._r.append(fldChar2)
            
        except Exception as e:
            print(f"添加页脚失败：{e}")
    
    def _setup_enhanced_styles(self):
        """设置增强版样式（支持中英文混合字体）GB/T 9704-2012"""
        from docx.enum.text import WD_LINE_SPACING
        # 获取或创建样式
        styles = self.document.styles

        # 1. 标题样式（7.3.1: 2号方正小标宋简体，居中，直接顶格无额外间距）
        title_style = self._create_or_get_style('Title_Style', WD_STYLE_TYPE.PARAGRAPH)
        self._set_style_font(title_style, '小标宋', self.english_font_name, Pt(22))
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_before = Pt(0)
        title_style.paragraph_format.space_after = Pt(0)
        title_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        title_style.paragraph_format.line_spacing = Pt(28.95)

        # 2. 一级标题样式（7.3.3: 第一层用黑体字，首行缩进2字符）
        h1_style = self._create_or_get_style('Heading1_Style', WD_STYLE_TYPE.PARAGRAPH)
        self._set_style_font(h1_style, '黑体', self.english_font_name, Pt(16))
        h1_style.paragraph_format.first_line_indent = Pt(32)  # 2字符缩进（16pt * 2 = 32pt）
        h1_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        h1_style.paragraph_format.space_before = Pt(0)
        h1_style.paragraph_format.space_after = Pt(0)
        h1_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        h1_style.paragraph_format.line_spacing = Pt(28.95)

        # 3. 二级标题样式（7.3.3: 第二层用楷体字，首行缩进2字符）
        h2_style = self._create_or_get_style('Heading2_Style', WD_STYLE_TYPE.PARAGRAPH)
        self._set_style_font(h2_style, '楷体', self.english_font_name, Pt(16))
        h2_style.paragraph_format.first_line_indent = Pt(32)  # 2字符缩进
        h2_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        h2_style.paragraph_format.space_before = Pt(0)
        h2_style.paragraph_format.space_after = Pt(0)
        h2_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        h2_style.paragraph_format.line_spacing = Pt(28.95)

        # 4. 正文样式（5.2.2: 3号仿宋_GB2312字; 两端对齐; 首行缩进2字符）
        normal_style = styles['Normal']
        self._set_style_font(normal_style, '仿宋_GB2312', self.english_font_name, Pt(16))
        normal_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        normal_style.paragraph_format.line_spacing = Pt(28.95)
        normal_style.paragraph_format.first_line_indent = Pt(32)  # 7.3.3: 左空二字（2字符）
        normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # 两端对齐

        # 5. 英文专用样式
        english_style = self._create_or_get_style('English_Style', WD_STYLE_TYPE.PARAGRAPH)
        english_style.font.name = self.english_font_name
        english_style.font.size = Pt(14)
        english_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        english_style.paragraph_format.line_spacing = Pt(28.95)
        english_style.paragraph_format.first_line_indent = Cm(0)

        # 6. 强调样式（加粗）
        strong_style = self._create_or_get_style('Strong_Style', WD_STYLE_TYPE.CHARACTER)
        self._set_style_font(strong_style, '黑体', self.english_font_name, Pt(16))
        strong_style.font.bold = True

        # 保存样式引用
        self.styles = {
            'title': title_style,
            'h1': h1_style,
            'h2': h2_style,
            'normal': normal_style,
            'english': english_style,
            'strong': strong_style
        }
    
    def _create_or_get_style(self, style_name: str, style_type: int):
        """创建或获取样式"""
        try:
            return self.document.styles[style_name]
        except KeyError:
            return self.document.styles.add_style(style_name, style_type)
    
    def _set_style_font(self, style, chinese_font: str, english_font: str, size: Pt):
        """设置样式字体（支持中英文分别设置）"""
        # 设置中文字体（尝试多个备选）
        font_candidates = self.chinese_font_mapping.get(chinese_font, [chinese_font])
        style.font.name = font_candidates[0]
        
        # 设置英文字体
        style._element.xpath('.//w:rFonts')[0].set(qn('w:ascii'), english_font)
        style._element.xpath('.//w:rFonts')[0].set(qn('w:hAnsi'), english_font)
        
        # 设置字号
        style.font.size = size
    
    def _generate_document_content(self, doc_type: str, config: Dict[str, Any]):
        """生成文档内容"""
        if doc_type == "meeting_minutes":
            self._generate_enhanced_meeting_minutes(config)
        elif doc_type == "speech":
            self._generate_enhanced_speech(config)
        elif doc_type == "discussion_outline":
            self._generate_enhanced_discussion_outline(config)
        elif doc_type == "work_report":
            self._generate_enhanced_work_report(config)
        elif doc_type == "notification":
            self._generate_notification(config)
        elif doc_type == "request_report":
            self._generate_request_report(config)
        else:
            raise ValueError(f"不支持的文档类型: {doc_type}")
    
    def _generate_enhanced_meeting_minutes(self, config: Dict[str, Any]):
        """生成增强版会议纪要"""
        # 标题
        title = config.get('title', '会议纪要')
        self._add_enhanced_paragraph(f"{title}会议纪要", style='title')
        
        # 会议基本情况
        self._add_heading("一、会议基本情况", level=1)
        
        meeting_info = config.get('meeting_info', {})
        self._add_enhanced_paragraph(f"（一）时间：{meeting_info.get('time', '')}")
        self._add_enhanced_paragraph(f"（二）地点：{meeting_info.get('location', '')}")
        self._add_enhanced_paragraph(f"（三）主持人：{meeting_info.get('host', '')}")
        self._add_enhanced_paragraph(f"（四）记录人：{meeting_info.get('recorder', '')}")
        
        # 参会人员
        attendees = meeting_info.get('attendees', [])
        if attendees:
            attendees_str = "、".join(attendees)
            self._add_enhanced_paragraph(f"（五）出席人员：{attendees_str}")
        
        absentees = meeting_info.get('absentees', [])
        if absentees:
            absentees_str = "、".join(absentees)
            self._add_enhanced_paragraph(f"（六）缺席人员：{absentees_str}")
        
        # 主要议题
        topics = config.get('topics', [])
        if topics:
            self._add_heading("二、会议议题", level=1)
            for i, topic in enumerate(topics, 1):
                self._add_enhanced_paragraph(f"{i}. {topic}")
        
        # 会议内容
        contents = config.get('contents', [])
        if contents:
            self._add_heading("三、会议主要内容", level=1)
            for i, content in enumerate(contents, 1):
                self._add_heading(f"（一）{content.get('topic', '')}", level=2)
                points = content.get('points', [])
                for j, point in enumerate(points, 1):
                    self._add_enhanced_paragraph(f"{j}. {point}")
        
        # 决议事项
        resolutions = config.get('resolutions', [])
        if resolutions:
            self._add_heading("四、会议决议事项", level=1)
            for i, resolution in enumerate(resolutions, 1):
                self._add_enhanced_paragraph(f"（一）{resolution.get('content', '')}")
                self._add_enhanced_paragraph(f"    责任部门/责任人：{resolution.get('responsible', '')}")
                self._add_enhanced_paragraph(f"    完成时限：{resolution.get('deadline', '')}")
        
        # 下一步工作安排
        next_steps = config.get('next_steps', [])
        if next_steps:
            self._add_heading("五、下一步工作安排", level=1)
            for i, step in enumerate(next_steps, 1):
                self._add_enhanced_paragraph(f"{i}. {step}")
        
        # 其他事项
        other_matters = config.get('other_matters', [])
        if other_matters:
            self._add_heading("六、其他事项", level=1)
            for i, matter in enumerate(other_matters, 1):
                self._add_enhanced_paragraph(f"{i}. {matter}")
        
        # 落款
        self._add_enhanced_paragraph("")
        self._add_enhanced_paragraph("")
        self._add_enhanced_paragraph("主持人签字：")
        self._add_enhanced_paragraph("记录人签字：")
        
        # 日期
        current_date = datetime.now().strftime("%Y年%m月%d日")
        self._add_enhanced_paragraph(f"{current_date}")
    
    def _generate_enhanced_speech(self, config: Dict[str, Any]):
        """生成增强版发言稿"""
        # 标题
        title = config.get('title', '发言稿')
        self._add_enhanced_paragraph(f"关于{title}的讲话", style='title')
        
        # 称呼
        self._add_enhanced_paragraph("")
        salutation = config.get('salutation', '尊敬的各位领导、同志们：')
        self._add_enhanced_paragraph(salutation)
        
        # 开场白
        introduction = config.get('introduction', '')
        if introduction:
            self._add_enhanced_paragraph("")
            self._add_enhanced_paragraph(introduction)
        
        # 正文
        sections = config.get('sections', [])
        for i, section in enumerate(sections, 1):
            self._add_enhanced_paragraph("")
            self._add_heading(f"一、{section.get('title', '')}", level=1)
            
            points = section.get('points', [])
            for j, point in enumerate(points, 1):
                self._add_heading(f"（一）{point.get('title', '')}", level=2)
                content = point.get('content', '')
                if content:
                    self._add_enhanced_paragraph(content)
        
        # 结束语
        conclusion = config.get('conclusion', '')
        if conclusion:
            self._add_enhanced_paragraph("")
            self._add_enhanced_paragraph(conclusion)
        
        # 结束语
        self._add_enhanced_paragraph("")
        self._add_enhanced_paragraph("谢谢大家！")
        
        # 落款
        self._add_enhanced_paragraph("")
        self._add_enhanced_paragraph("")
        organization = config.get('organization', '')
        if organization:
            self._add_enhanced_paragraph(organization)
        
        position = config.get('position', '')
        if position:
            self._add_enhanced_paragraph(position)
        
        name = config.get('name', '')
        if name:
            self._add_enhanced_paragraph(name)
        
        # 日期
        current_date = datetime.now().strftime("%Y年%m月%d日")
        self._add_enhanced_paragraph(current_date)
    
    def _generate_enhanced_discussion_outline(self, config: Dict[str, Any]):
        """生成增强版讨论提纲"""
        # 标题
        title = config.get('title', '讨论提纲')
        self._add_enhanced_paragraph(f"{title}讨论提纲", style='title')
        
        # 讨论背景
        background = config.get('background', '')
        if background:
            self._add_heading("一、讨论背景", level=1)
            self._add_enhanced_paragraph(background)
        
        # 讨论目标
        objectives = config.get('objectives', [])
        if objectives:
            self._add_heading("二、讨论目标", level=1)
            for i, objective in enumerate(objectives, 1):
                self._add_enhanced_paragraph(f"{i}. {objective}")
        
        # 讨论要点
        points = config.get('points', [])
        if points:
            self._add_heading("三、讨论要点", level=1)
            for i, point in enumerate(points, 1):
                self._add_heading(f"（一）{point.get('title', '')}", level=2)
                subpoints = point.get('subpoints', [])
                for j, subpoint in enumerate(subpoints, 1):
                    self._add_enhanced_paragraph(f"{j}. {subpoint}")
        
        # 预期成果
        expected_outcomes = config.get('expected_outcomes', [])
        if expected_outcomes:
            self._add_heading("四、预期成果", level=1)
            for i, outcome in enumerate(expected_outcomes, 1):
                self._add_enhanced_paragraph(f"{i}. {outcome}")
        
        # 时间安排
        schedule = config.get('schedule', [])
        if schedule:
            self._add_heading("五、时间安排", level=1)
            for i, item in enumerate(schedule, 1):
                self._add_enhanced_paragraph(f"{i}. {item}")
        
        # 参考材料
        references = config.get('references', [])
        if references:
            self._add_heading("六、参考材料", level=1)
            for i, reference in enumerate(references, 1):
                self._add_enhanced_paragraph(f"{i}. {reference}")
    
    def _generate_enhanced_work_report(self, config: Dict[str, Any]):
        """生成增强版工作汇报"""
        # 标题
        title = config.get('title', '工作汇报')
        self._add_enhanced_paragraph(f"关于{title}的工作汇报", style='title')
        
        # 主送机关
        recipient = config.get('recipient', '')
        if recipient:
            self._add_enhanced_paragraph("")
            self._add_enhanced_paragraph(f"{recipient}：")
        
        # 工作进展情况
        progress = config.get('progress', [])
        if progress:
            self._add_heading("一、工作进展情况", level=1)
            for i, item in enumerate(progress, 1):
                self._add_heading(f"（一）{item.get('area', '')}", level=2)
                details = item.get('details', [])
                for j, detail in enumerate(details, 1):
                    self._add_enhanced_paragraph(f"{j}. {detail}")
        
        # 主要成绩和亮点
        achievements = config.get('achievements', [])
        if achievements:
            self._add_heading("二、主要成绩和亮点", level=1)
            for i, achievement in enumerate(achievements, 1):
                self._add_enhanced_paragraph(f"{i}. {achievement}")
        
        # 存在的问题和困难
        problems = config.get('problems', [])
        if problems:
            self._add_heading("三、存在的问题和困难", level=1)
            for i, problem in enumerate(problems, 1):
                self._add_enhanced_paragraph(f"{i}. {problem}")
        
        # 下一步工作计划
        plans = config.get('plans', [])
        if plans:
            self._add_heading("四、下一步工作计划", level=1)
            for i, plan in enumerate(plans, 1):
                self._add_enhanced_paragraph(f"{i}. {plan}")
        
        # 需要支持的事项
        support_needed = config.get('support_needed', [])
        if support_needed:
            self._add_heading("五、需要支持的事项", level=1)
            for i, support in enumerate(support_needed, 1):
                self._add_enhanced_paragraph(f"{i}. {support}")
        
        # 结束语
        self._add_enhanced_paragraph("")
        self._add_enhanced_paragraph("特此汇报。")
        
        # 落款
        self._add_enhanced_paragraph("")
        self._add_enhanced_paragraph("")
        organization = config.get('organization', '')
        if organization:
            self._add_enhanced_paragraph(organization)
        
        # 日期
        current_date = datetime.now().strftime("%Y年%m月%d日")
        self._add_enhanced_paragraph(current_date)
    
    def _generate_notification(self, config: Dict[str, Any]):
        """生成通知"""
        # 标题
        title = config.get('title', '通知')
        self._add_enhanced_paragraph(f"关于{title}的通知", style='title')
        
        # 主送机关
        recipient = config.get('recipient', '各有关单位：')
        self._add_enhanced_paragraph("")
        self._add_enhanced_paragraph(recipient)
        
        # 正文
        content = config.get('content', '')
        if content:
            self._add_enhanced_paragraph("")
            self._add_enhanced_paragraph(content)
        
        # 具体要求
        requirements = config.get('requirements', [])
        if requirements:
            self._add_enhanced_paragraph("")
            for i, req in enumerate(requirements, 1):
                self._add_enhanced_paragraph(f"{i}. {req}")
        
        # 落款
        self._add_enhanced_paragraph("")
        self._add_enhanced_paragraph("")
        organization = config.get('organization', '')
        if organization:
            self._add_enhanced_paragraph(organization)
        
        # 日期
        current_date = datetime.now().strftime("%Y年%m月%d日")
        self._add_enhanced_paragraph(current_date)
    
    def _generate_request_report(self, config: Dict[str, Any]):
        """生成请示报告"""
        # 标题
        title = config.get('title', '请示')
        self._add_enhanced_paragraph(f"关于{title}的请示", style='title')
        
        # 主送机关
        recipient = config.get('recipient', '')
        if recipient:
            self._add_enhanced_paragraph("")
            self._add_enhanced_paragraph(f"{recipient}：")
        
        # 请示事项
        request_items = config.get('request_items', [])
        if request_items:
            self._add_enhanced_paragraph("")
            for i, item in enumerate(request_items, 1):
                self._add_enhanced_paragraph(f"一、{item.get('title', '')}")
                content = item.get('content', '')
                if content:
                    self._add_enhanced_paragraph(content)
        
        # 请示理由
        reasons = config.get('reasons', [])
        if reasons:
            self._add_enhanced_paragraph("")
            self._add_heading("二、请示理由", level=1)
            for i, reason in enumerate(reasons, 1):
                self._add_enhanced_paragraph(f"{i}. {reason}")
        
        # 建议方案
        suggestions = config.get('suggestions', [])
        if suggestions:
            self._add_enhanced_paragraph("")
            self._add_heading("三、建议方案", level=1)
            for i, suggestion in enumerate(suggestions, 1):
                self._add_enhanced_paragraph(f"{i}. {suggestion}")
        
        # 结束语
        self._add_enhanced_paragraph("")
        self._add_enhanced_paragraph("以上请示妥否，请批示。")
        
        # 落款
        self._add_enhanced_paragraph("")
        self._add_enhanced_paragraph("")
        organization = config.get('organization', '')
        if organization:
            self._add_enhanced_paragraph(organization)
        
        # 日期
        current_date = datetime.now().strftime("%Y年%m月%d日")
        self._add_enhanced_paragraph(current_date)
    
    def _add_heading(self, text: str, level: int = 1):
        """添加标题"""
        if level == 1:
            self._add_enhanced_paragraph(text, style='h1')
        elif level == 2:
            self._add_enhanced_paragraph(text, style='h2')
    
    def _add_enhanced_paragraph(self, text: str, style: str = 'normal'):
        """添加增强版段落（支持中英文混合处理）"""
        if not self.document:
            return
        
        # 创建段落
        p = self.document.add_paragraph()
        
        # 设置段落样式
        if style in self.styles:
            p.style = self.styles[style]
        
        # 处理文本中的中英文混合内容
        if text:
            self._add_mixed_language_run(p, text, style)
    
    def _add_mixed_language_run(self, paragraph, text: str, style: str):
        """添加混合语言文本（中英文分别处理）"""
        # 使用正则表达式分割中英文
        pattern = r'([a-zA-Z0-9\s\.\,\!\?\;\:\'\"\(\)\[\]\{\}\-\+\=\*\/\\\&\#\@\%\$\^\|\<\>]+)'
        parts = re.split(pattern, text)
        
        for part in parts:
            if not part:
                continue
            
            # 判断是否为英文（包含字母或数字）
            if re.search(r'[a-zA-Z0-9]', part):
                # 英文部分使用英文样式或特殊处理
                run = paragraph.add_run(part)
                run.font.name = self.english_font_name
                
                # 根据样式调整字号
                if style == 'normal':
                    run.font.size = Pt(14)  # 英文比中文小2磅
                else:
                    # 其他样式保持与中文相同字号
                    run.font.size = self.styles[style].font.size
            else:
                # 中文部分使用段落样式
                run = paragraph.add_run(part)
    
    def _summarize_config(self, config: Dict[str, Any]) -> Dict[str, Any]:
        """汇总配置信息（用于修订历史）"""
        summary = {
            'config_keys': list(config.keys()),
            'has_meeting_info': 'meeting_info' in config,
            'has_sections': 'sections' in config,
            'has_topics': 'topics' in config,
            'config_size': len(str(config))
        }
        
        # 添加部分关键信息（避免敏感信息）
        if 'title' in config:
            summary['title'] = config['title']
        
        return summary
    
    def save_document(self, output_path: str, author: str = "系统生成"):
        """保存文档并记录修订历史"""
        if not self.document:
            print("错误：没有文档可保存")
            return False
        
        try:
            # 保存文档
            self.document.save(output_path)
            print(f"✅ 文档已保存到：{output_path}")
            
            # 记录文档保存（如果启用了历史跟踪）
            if self.enable_history_tracking and self.revision_tracker:
                # 读取文档内容计算哈希
                with open(output_path, 'rb') as f:
                    content = f.read()
                
                # 创建修改记录
                changes = [{
                    'type': 'save',
                    'description': '文档保存',
                    'file_size': len(content),
                    'output_path': output_path
                }]
                
                revision_record = self.revision_tracker.track_document_modification(
                    document_path=output_path,
                    document_type='unknown',  # 文档类型未知
                    author=author,
                    changes=changes,
                    metadata={'action': 'save', 'generator': 'enhanced'}
                )
                
                print(f"📝 文档保存记录：{revision_record['id']}")
                print(f"📁 修订历史保存到：{revision_record.get('history_file', '未知')}")
            
            return True
            
        except Exception as e:
            print(f"❌ 保存文档失败：{e}")
            return False
    
    def get_revision_summary(self, document_path: str) -> Optional[Dict[str, Any]]:
        """获取文档修订摘要"""
        if not self.revision_tracker:
            return None
        
        return self.revision_tracker.get_document_summary(document_path)


def load_config(config_file: str) -> Dict[str, Any]:
    """加载配置文件"""
    if not os.path.exists(config_file):
        print(f"错误：配置文件不存在：{config_file}")
        sys.exit(1)
    
    try:
        with open(config_file, 'r', encoding='utf-8') as f:
            if config_file.endswith('.json'):
                return json.load(f)
            else:
                print(f"错误：不支持的配置文件格式：{config_file}")
                sys.exit(1)
    except Exception as e:
        print(f"错误：加载配置文件失败：{e}")
        sys.exit(1)


def create_example_config(doc_type: str) -> Dict[str, Any]:
    """创建示例配置"""
    if doc_type == "meeting_minutes":
        return {
            "title": "数字化转型工作",
            "meeting_info": {
                "time": "2026年3月28日 14:00-16:00",
                "location": "第一会议室",
                "host": "张三",
                "recorder": "李四",
                "attendees": ["王五", "赵六", "孙七"],
                "absentees": ["周八"]
            },
            "topics": [
                "数字化转型工作方案讨论",
                "技术平台选型 (Technology Platform Selection)",
                "实施计划安排"
            ],
            "contents": [
                {
                    "topic": "数字化转型工作方案讨论",
                    "points": [
                        "分析了当前数字化发展形势",
                        "讨论了数字化转型的必要性 (The necessity of digital transformation)",
                        "明确了数字化转型的目标"
                    ]
                }
            ],
            "resolutions": [
                {
                    "content": "原则通过《数字化转型工作方案》",
                    "responsible": "技术部 (Technology Department)",
                    "deadline": "2026年4月15日前"
                }
            ],
            "next_steps": [
                "技术部负责完善方案细节",
                "人力资源部负责人员培训安排 (HR Department arranges training)",
                "财务部负责预算编制"
            ]
        }
    elif doc_type == "speech":
        return {
            "title": "推进数字化转型工作",
            "salutation": "尊敬的各位领导、同志们：",
            "introduction": "今天我们召开这次会议，主要目的是研究部署推进数字化转型工作。下面，我讲几点意见：",
            "sections": [
                {
                    "title": "充分认识数字化转型的重要意义",
                    "points": [
                        {
                            "title": "数字化转型是时代发展的必然要求",
                            "content": "当前，数字化浪潮席卷全球，新一代信息技术迅猛发展，数字化转型已成为企业发展的关键路径。Digital transformation is a key path for enterprise development."
                        }
                    ]
                }
            ],
            "conclusion": "同志们，数字化转型任务艰巨、责任重大。让我们以更加饱满的热情、更加务实的作风，扎实推进各项工作，确保数字化转型取得实效。",
            "organization": "某某单位",
            "position": "主任 (Director)",
            "name": "张三"
        }
    else:
        return {}


def main():
    """命令行接口"""
    parser = argparse.ArgumentParser(description='增强版官方文档生成器')
    parser.add_argument('--type', required=True, 
                       choices=['meeting_minutes', 'speech', 'discussion_outline', 
                               'work_report', 'notification', 'request_report'],
                       help='文档类型')
    parser.add_argument('--config', help='配置文件路径（JSON格式）')
    parser.add_argument('--output', required=True, help='输出文件路径')
    parser.add_argument('--author', default='系统生成', help='作者/操作者')
    parser.add_argument('--example', action='store_true', help='使用示例配置')
    parser.add_argument('--no-history', action='store_true', help='禁用修订历史跟踪')
    parser.add_argument('--show-summary', action='store_true', help='显示修订摘要')
    
    args = parser.parse_args()
    
    # 加载配置
    if args.example:
        config = create_example_config(args.type)
    elif args.config:
        config = load_config(args.config)
    else:
        print("错误：必须提供配置文件或使用--example参数")
        parser.print_help()
        sys.exit(1)
    
    # 生成文档
    try:
        generator = EnhancedDocumentGenerator(enable_history_tracking=not args.no_history)
        generator.create_document(args.type, config, args.author)
        
        # 保存文档
        if generator.save_document(args.output, args.author):
            print("\n" + "="*60)
            print("✅ 文档生成成功！")
            print(f"📄 文档类型：{args.type}")
            print(f"📁 输出路径：{args.output}")
            print(f"👤 作者：{args.author}")
            
            # 显示修订摘要
            if not args.no_history and args.show_summary:
                summary = generator.get_revision_summary(args.output)
                if summary:
                    print(f"\n📊 修订摘要：")
                    print(f"   修订次数：{summary.get('revision_count', 0)}")
                    print(f"   首次修订：{summary.get('first_revision', '无')}")
                    print(f"   最后修订：{summary.get('latest_revision', '无')}")
                    print(f"   最后操作：{summary.get('latest_action', '无')}")
                    print(f"   最后作者：{summary.get('latest_author', '无')}")
            
            print("="*60)
            
    except Exception as e:
        print(f"❌ 错误：生成文档失败：{e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()