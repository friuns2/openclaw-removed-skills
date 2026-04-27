#!/usr/bin/env python3
"""
generate_agenda.py - 会议议程生成脚本
根据传入的结构化数据，生成 .docx 和 .pdf 格式的会议议程文件。
PDF 由 Word 文档转换生成，确保完全一致。

依赖：
    pip install python-docx docx2pdf

用法（由 WorkBuddy 调用）：
    python generate_agenda.py --data agenda.json --output ./output
"""

import argparse
import json
import os
import sys
from pathlib import Path


# ── 依赖检查 ────────────────────────────────────────────────────────────────────
def ensure_deps():
    import importlib.util
    missing = []
    for pkg, import_name in [("python-docx", "docx"), ("docx2pdf", "docx2pdf")]:
        if importlib.util.find_spec(import_name) is None:
            missing.append(pkg)
    if missing:
        import subprocess
        subprocess.check_call(
            [sys.executable, "-m", "pip", "install", "--quiet"] + missing
        )


ensure_deps()

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── 字体常量 ─────────────────────────────────────────────────────────────────────
FONT_HEITI    = "黑体"
FONT_FANGSONG = "仿宋_GB2312"
FONT_KAITI    = "楷体_GB2312"

# ── 中文数字 ─────────────────────────────────────────────────────────────────────
CN_NUMS = "一二三四五六七八九十"

def cn_num(n):
    return CN_NUMS[n - 1] if 1 <= n <= 10 else str(n)


# ── 底层工具函数 ─────────────────────────────────────────────────────────────────

def set_run_font(run, cn_font, en_font="Times New Roman", size_pt=16, bold=False):
    """设置 run 的中英文字体、字号、加粗"""
    run.font.name = en_font
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), cn_font)
    rFonts.set(qn("w:hint"), "eastAsia")
    rPr.insert(0, rFonts)


def set_para_spacing(para, before=0, after=0):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"),  str(after))
    pPr.append(spacing)


# ── 段落构建函数 ─────────────────────────────────────────────────────────────────

def add_title(doc, text):
    """
    会议议程标题：仿宋_GB2312 小一（24pt）加粗 居中
    """
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(para, before=0, after=240)
    run = para.add_run(text)
    set_run_font(run, FONT_FANGSONG, size_pt=24, bold=True)
    return para


def add_section_heading(doc, text):
    """
    一级标题（环节 / 人员）：黑体 三号（16pt）加粗 左对齐
    格式：一、XX
    """
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para_spacing(para, before=200, after=80)
    run = para.add_run(text)
    set_run_font(run, FONT_HEITI, size_pt=16, bold=True)
    return para


def add_field_line(doc, label, value, size_pt=16):
    """
    字段行：标签黑体加粗，值仿宋
    格式：时间：XXXX
    """
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para_spacing(para, before=60, after=60)
    run_label = para.add_run(f"{label}：")
    set_run_font(run_label, FONT_HEITI, size_pt=size_pt, bold=True)
    run_val = para.add_run(str(value))
    set_run_font(run_val, FONT_FANGSONG, size_pt=size_pt)
    return para


def add_agenda_items(doc, items, size_pt=16):
    """议程列表：仿宋 三号，带序号"""
    for i, item in enumerate(items, 1):
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_para_spacing(para, before=40, after=40)
        run = para.add_run(f"  {i}. {item}")
        set_run_font(run, FONT_FANGSONG, size_pt=size_pt)


def add_persons_section(doc, units, size_pt=16):
    """
    参会人员，按单位分组
    - 单位名称：楷体_GB2312 三号 加粗，前空两个全角空格
    - 成员：仿宋_GB2312 三号，不加"-"，姓名与职务之间用两个全角空格
    """
    for unit in units:
        unit_name = unit.get("unit", "")
        members   = unit.get("members", [])

        # 单位标题：前空两个全角空格，楷体加粗
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_para_spacing(para, before=120, after=60)
        run = para.add_run(f"\u3000\u3000{unit_name}")
        set_run_font(run, FONT_KAITI, size_pt=size_pt, bold=True)

        # 成员列表：无"-"，姓名＋全角空格＋职务
        for member in members:
            name  = member.get("name", "")
            title = member.get("title", "")
            note  = member.get("note", "")
            para_m = doc.add_paragraph()
            para_m.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_para_spacing(para_m, before=30, after=30)
            line = f"  {name}\u3000\u3000{title}"
            run_m = para_m.add_run(line)
            set_run_font(run_m, FONT_FANGSONG, size_pt=size_pt)
            # 补充说明另起一行
            if note:
                para_note = doc.add_paragraph()
                para_note.alignment = WD_ALIGN_PARAGRAPH.LEFT
                set_para_spacing(para_note, before=0, after=20)
                run_note = para_note.add_run(f"    {note}")
                set_run_font(run_note, FONT_FANGSONG, size_pt=size_pt)


# ── Word 文档构建 ────────────────────────────────────────────────────────────────

def build_docx(data: dict, output_path: str):
    doc = Document()

    # 页面设置 A4
    sec = doc.sections[0]
    sec.page_width    = Cm(21)
    sec.page_height   = Cm(29.7)
    sec.left_margin   = Cm(3.18)
    sec.right_margin  = Cm(3.18)
    sec.top_margin    = Cm(2.54)
    sec.bottom_margin = Cm(2.54)

    doc.styles["Normal"].font.size = Pt(16)

    # ── 会议标题（仿宋_GB2312 小一 24pt 加粗 居中）──
    title = data.get("title", "会议议程")
    add_title(doc, title)

    # ── 环节计数（人员环节也占序号）──
    sections_data = data.get("sections", [])
    section_counter = 0  # 滚动序号

    for sec_data in sections_data:
        sec_name = sec_data.get("name", "")
        section_counter += 1

        # 一级标题
        add_section_heading(doc, f"{cn_num(section_counter)}、{sec_name}")

        if sec_data.get("time"):
            add_field_line(doc, "时间", sec_data["time"])
        if sec_data.get("location"):
            add_field_line(doc, "地点", sec_data["location"])
        if sec_data.get("content"):
            add_field_line(doc, "内容", sec_data["content"])
        if sec_data.get("host"):
            add_field_line(doc, "主持人", sec_data["host"])

        # 议程
        if sec_data.get("agenda"):
            para_ag = doc.add_paragraph()
            set_para_spacing(para_ag, before=60, after=40)
            set_run_font(para_ag.add_run("议程："), FONT_HEITI, size_pt=16, bold=True)
            add_agenda_items(doc, sec_data["agenda"])

        # 人员 —— 作为独立一级标题环节
        if sec_data.get("persons"):
            section_counter += 1
            add_section_heading(doc, f"{cn_num(section_counter)}、人员")
            add_persons_section(doc, sec_data["persons"])

    doc.save(output_path)
    print(f"[✓] Word 文档已保存：{output_path}")
    return output_path


# ── PDF：由 Word 转换（确保与 Word 完全一致）──────────────────────────────────────

def build_pdf_from_docx(docx_path: str, pdf_path: str):
    """
    将 docx 转换为 PDF，依次尝试以下方案：
      1. win32com（Windows + Word 已安装）
      2. LibreOffice 命令行
      3. docx2pdf 库
    任何方案成功即停止。全部失败则提示手动转换。
    """
    docx_abs = str(Path(docx_path).resolve())
    pdf_abs  = str(Path(pdf_path).resolve())
    pdf_dir  = str(Path(pdf_path).resolve().parent)

    # ── 方案 1：win32com + pythoncom（修复 COM 线程模型问题）──
    try:
        import win32com.client
        import pythoncom
        pythoncom.CoInitialize()
        try:
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            doc = word.Documents.Open(docx_abs)
            doc.SaveAs(pdf_abs, FileFormat=17)  # 17 = wdFormatPDF
            doc.Close()
            word.Quit()
            print(f"[✓] PDF 文档已保存（via Word COM）：{pdf_abs}")
            return
        finally:
            pythoncom.CoUninitialize()
    except ImportError:
        pass  # win32com 未安装，跳过
    except Exception as e:
        print(f"[i] Word COM 转换失败（{e}），尝试下一方案…", file=sys.stderr)

    # ── 方案 2：LibreOffice 命令行 ──
    import shutil, subprocess
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    for candidate in [
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ]:
        if os.path.exists(candidate):
            soffice = candidate
            break
    if soffice:
        try:
            result = subprocess.run(
                [soffice, "--headless", "--convert-to", "pdf",
                 "--outdir", pdf_dir, docx_abs],
                capture_output=True, timeout=60
            )
            # LibreOffice 输出文件名与 docx 同名
            lo_pdf = str(Path(pdf_dir) / (Path(docx_abs).stem + ".pdf"))
            if os.path.exists(lo_pdf) and lo_pdf != pdf_abs:
                os.replace(lo_pdf, pdf_abs)
            if os.path.exists(pdf_abs):
                print(f"[✓] PDF 文档已保存（via LibreOffice）：{pdf_abs}")
                return
        except Exception as e:
            print(f"[i] LibreOffice 转换失败（{e}），尝试下一方案…", file=sys.stderr)

    # ── 方案 3：docx2pdf 库 ──
    try:
        from docx2pdf import convert
        convert(docx_abs, pdf_abs)
        print(f"[✓] PDF 文档已保存（via docx2pdf）：{pdf_abs}")
        return
    except Exception as e:
        print(f"[i] docx2pdf 转换失败（{e}）", file=sys.stderr)

    # ── 全部失败 ──
    print(
        "[!] PDF 自动转换未成功。\n"
        "    请在 Microsoft Word 或 WPS 中打开 .docx 文件，选择「另存为 PDF」手动导出。",
        file=sys.stderr,
    )


# ── 主入口 ───────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="生成会议议程 Word + PDF")
    parser.add_argument("--data",   required=True, help="JSON 数据文件路径")
    parser.add_argument("--output", default=".",   help="输出目录")
    args = parser.parse_args()

    with open(args.data, "r", encoding="utf-8") as f:
        data = json.load(f)

    out_dir = Path(args.output)
    out_dir.mkdir(parents=True, exist_ok=True)

    safe_name = data.get("filename", "会议议程")
    docx_path = str(out_dir / f"{safe_name}.docx")
    pdf_path  = str(out_dir / f"{safe_name}.pdf")

    build_docx(data, docx_path)
    build_pdf_from_docx(docx_path, pdf_path)

    print(f"\n[完成] 文件已生成：\n  Word: {docx_path}\n  PDF:  {pdf_path}")


if __name__ == "__main__":
    main()
